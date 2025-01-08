from flask import Flask, render_template, redirect, url_for, request, session, flash, jsonify, send_file
from flask_pymongo import PyMongo
from werkzeug.security import generate_password_hash, check_password_hash
from client_features import get_clients, get_client_by_id, update_client, delete_client
from docx import Document
from bson import ObjectId
import io
import logging
import openpyxl
import pandas as pd
from openpyxl.styles import Font, Alignment
from openpyxl.utils import get_column_letter
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.worksheet.table import Table, TableStyleInfo
from io import BytesIO
from openpyxl import Workbook
from flask import Flask, request, redirect, url_for, send_from_directory, render_template, flash
from werkzeug.utils import secure_filename
from appwrite.client import Client
from appwrite.services.storage import Storage
from appwrite.services.databases import Databases
from appwrite.exception import AppwriteException
from dotenv import load_dotenv
import os
import requests
import uuid
import mimetypes
from docx.shared import Pt
from pymongo.errors import WriteConcernError
import gridfs
import time
from bson import json_util
import json
from datetime import datetime, timedelta
from urllib.parse import quote
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl.styles import Font, Alignment, PatternFill
from openpyxl import Workbook
from werkzeug.exceptions import HTTPException
from flask import json
import traceback
import logging
from werkzeug.security import generate_password_hash
from itsdangerous import URLSafeTimedSerializer
from flask_mail import Mail, Message
import os
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import logging
from jinja2 import Environment, FileSystemLoader
from openpyxl.styles import PatternFill, Border, Side, Alignment, Font
from openpyxl.utils import get_column_letter
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document
import io
from datetime import datetime
from flask import send_file
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import io
from datetime import datetime
from urllib.parse import unquote
from user_agents import parse
from financials import financials




# Load environment variables
load_dotenv()

app = Flask(__name__)
app.secret_key = os.getenv("SECRET_KEY")





# MongoDB configuration
app.config['MONGO_URI'] = os.getenv("MONGO_URI")
mongo = PyMongo(app)


# Configure Appwrite Client
client = Client()
client.set_endpoint(os.getenv("APPWRITE_ENDPOINT"))  # Your Appwrite server URL
client.set_project(os.getenv("APPWRITE_PROJECT_ID"))  # Your project ID
client.set_key(os.getenv("APPWRITE_API_KEY"))  # Your secret API key

# Initialize Appwrite Storage and Database IDs
storage = Storage(client)
databases = Databases(client)
database_id = os.getenv("APPWRITE_DATABASE_ID")
collection_id = os.getenv("APPWRITE_COLLECTION_ID")
bucket_id = os.getenv("APPWRITE_BUCKET_ID")


APPWRITE_ENDPOINT = os.getenv("APPWRITE_ENDPOINT")
APPWRITE_PROJECT_ID = os.getenv("APPWRITE_PROJECT_ID")
APPWRITE_API_KEY = os.getenv("APPWRITE_API_KEY")
BUCKET_ID = os.getenv("APPWRITE_BUCKET_ID")


                  # Replace with your actual bucket ID

logging.basicConfig(level=logging.DEBUG, filename="/tmp/app.log", format="%(asctime)s - %(levelname)s - %(message)s")

# Register the blueprint
app.register_blueprint(financials)

# Allowed file extensions
ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx', 'xls', 'xlsx', 'txt', 'csv'}


# Access the database via mongo.db
db = mongo.db
fs = gridfs.GridFS(db)

# Defining collections globally
clients_collection = db["clients"]
audit_scope_collection = db["audit_scope"]
client_audit_plan_collection = db["Client_Audit_Plan"]
team_users_collection = db["team_users"]
audit_execution_collection=db["audit_execution_collection"]
generate_audit_procedures_collection = db["generate_audit_procedures"]
generate_audit_procedures_collection = db['generate_audit_procedures_collection']
modified_audit_procedures_collection = db['Modified_Audit_Procedures']
inventory_management_collection = db['Audit_procedures_inventory_management']
audit_execution_dynamic = db["audit_execution_dynamic"] 
Audit_procedures_production_operations = db['Audit_procedures_production_operations']

# Additional collections for visitor tracking
visitors_collection = db["visitors"]
visitor_analytics_collection = db["visitor_analytics"]
cookie_consents_collection = db["cookie_consents"]


db.audit_execution_dynamic.create_index([("client_id", 1)])
db.audit_execution_dynamic.create_index([("tasks.updated_at", -1)])
db.client_audit_plan_collection.create_index([("client_id", 1), ("start_date", 1)])
db.users_collection.create_index([("client_id", 1)])

# Create necessary indexes
db.audit_execution_dynamic.create_index([("client_id", 1)])
db.audit_execution_dynamic.create_index([("period_start", 1), ("period_end", 1)])
db.audit_execution_dynamic.create_index([("client_id", 1), ("period_start", 1), ("period_end", 1)])

UPLOAD_FOLDER = '/tmp/uploads' # Ensure this folder exists in your project directory
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Ensure the uploads folder exists
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# Set up logging
logging.basicConfig(level=logging.DEBUG, filename="audit_app.log",
                    format="%(asctime)s - %(levelname)s - %(message)s")

# Initialize Flask-Mail
mail = Mail(app)

# Configure email settings
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 587
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USERNAME'] = os.getenv('rockstarpraveen44@gmail.com')
app.config['MAIL_PASSWORD'] = os.getenv('S.pr@veen225')
app.config['MAIL_DEFAULT_SENDER'] = os.getenv('AI_BY_CA')

# Initialize serializer for token generation
serializer = URLSafeTimedSerializer(app.config['SECRET_KEY'])


@app.route('/')
def home():
    return render_template('landing.html')

@app.route('/about')
def about():
    return render_template('about.html')

@app.route('/contact', methods=['GET', 'POST'])
def contact():
    try:
        if request.method == 'POST':
            # Get form data
            message_data = {
                "name": request.form.get('name'),
                "email": request.form.get('email'),
                "subject": request.form.get('subject'),
                "message": request.form.get('message'),
                "status": "unread",  # Track message status
                "created_at": datetime.now(),
                "updated_at": datetime.now()
            }

            try:
                # Insert into contact_messages collection
                result = mongo.db.contact_messages.insert_one(message_data)
                
                if result.inserted_id:
                    # Log successful submission
                    logging.info(f"Contact message saved successfully from {message_data['email']}")
                    flash('Your message has been sent successfully! We will get back to you soon.', 'success')
                else:
                    logging.error("Failed to save contact message")
                    flash('Sorry, there was an error sending your message. Please try again.', 'error')
                
                return redirect(url_for('contact'))

            except Exception as e:
                logging.error(f"Database error: {str(e)}")
                flash('An error occurred while saving your message. Please try again.', 'error')
                return redirect(url_for('contact'))

        # GET request - show contact form
        return render_template('contact.html')

    except Exception as e:
        logging.error(f"Unexpected error in contact route: {str(e)}")
        flash('An unexpected error occurred. Please try again later.', 'error')
        return redirect(url_for('home'))



@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        
        user = mongo.db.users.find_one({"username": username})
        
        if user and check_password_hash(user['password'], password):
            session['logged_in'] = True
            session['username'] = username
            session['user_id'] = str(user['_id'])
            session['role'] = user.get('role', 'user')
            
            flash('Login successful!', 'success')
            return redirect(url_for('dashboard'))
        else:
            flash('Invalid username or password', 'error')
    
    return render_template('login.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        try:
            print("Received registration request") # Debug log
            
            # Get form data
            username = request.form.get('username')
            email = request.form.get('email')
            mobile = request.form.get('mobile')
            password = request.form.get('password')
            role = request.form.get('role')

            # Validate required fields
            if not all([username, email, mobile, password, role]):
                return jsonify({
                    'success': False,
                    'message': 'All fields are required'
                }), 400

            # Check if user already exists
            existing_user = mongo.db.users.find_one({
                "$or": [
                    {"username": username},
                    {"email": email},
                    {"mobile": mobile}
                ]
            })

            if existing_user:
                if existing_user.get('username') == username:
                    message = "Username already exists"
                elif existing_user.get('email') == email:
                    message = "Email already registered"
                else:
                    message = "Mobile number already registered"
                    
                return jsonify({
                    'success': False,
                    'message': message
                }), 400

            # Create new user
            hashed_password = generate_password_hash(password)
            new_user = {
                "username": username,
                "email": email,
                "mobile": mobile,
                "password": hashed_password,
                "role": role,
                "created_at": datetime.utcnow()
            }

            # Insert into database
            mongo.db.users.insert_one(new_user)

            print("User registered successfully") # Debug log
            return jsonify({
                'success': True,
                'message': 'Registration successful! Please login.'
            })

        except Exception as e:
            print(f"Registration error: {str(e)}") # Debug log
            return jsonify({
                'success': False,
                'message': f'An error occurred: {str(e)}'
            }), 500

    return render_template('register.html')

def send_reset_email(email, reset_url):
    sender_email = "rockstarpraveen44@gmail.com"
    sender_password = "uqmt yfwc khzo gkgj"  # Your app password

    # Create message
    msg = MIMEMultipart()
    msg['From'] = 'Audit IQ AI <{}>'.format(sender_email)
    msg['To'] = email
    msg['Subject'] = "Password Reset Request"

    # Create HTML body
    html = f"""
    <html>
        <body style="font-family: Arial, sans-serif; line-height: 1.6; color: #333;">
            <div style="max-width: 600px; margin: 0 auto; padding: 20px;">
                <div style="background: #4361ee; color: white; padding: 20px; text-align: center; border-radius: 5px;">
                    <h2>Password Reset Request</h2>
                </div>
                <div style="padding: 20px; background: #f8f9fa; border-radius: 5px; margin-top: 20px;">
                    <p>Hello,</p>
                    <p>We received a request to reset your password for your Audit IQ AI account.</p>
                    <p>Click the button below to reset your password:</p>
                    <div style="text-align: center; margin: 30px 0;">
                        <a href="{reset_url}" 
                           style="background: #4361ee; color: white; padding: 12px 25px; text-decoration: none; border-radius: 5px; display: inline-block;">
                            Reset Password
                        </a>
                    </div>
                    <p>If you didn't request this, you can safely ignore this email.</p>
                    <p>This link will expire in 1 hour for security reasons.</p>
                    <p>Best regards,<br>Audit IQ AI Team</p>
                </div>
            </div>
        </body>
    </html>
    """

    msg.attach(MIMEText(html, 'html'))

    try:
        # Create SMTP session
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        
        # Login
        server.login(sender_email, sender_password)
        
        # Send email
        server.send_message(msg)
        
        # Close connection
        server.quit()
        return True
        
    except Exception as e:
        print(f"Error sending email: {str(e)}")
        return False
    
    
# Add these routes to your app.py

@app.route('/reset-password/<token>', methods=['GET', 'POST'])
def reset_password(token):
    try:
        # Verify the token
        email = verify_reset_token(token)
        if not email:
            flash('Invalid or expired reset link', 'error')
            return redirect(url_for('forgot_password'))

        if request.method == 'POST':
            password = request.form.get('password')
            confirm_password = request.form.get('confirm_password')

            if not password or not confirm_password:
                flash('Please enter both password fields', 'error')
                return redirect(url_for('reset_password', token=token))

            if password != confirm_password:
                flash('Passwords do not match', 'error')
                return redirect(url_for('reset_password', token=token))

            # Update password in database
            hashed_password = generate_password_hash(password)
            mongo.db.users.update_one(
                {"email": email},
                {"$set": {"password": hashed_password}}
            )

            flash('Your password has been updated!', 'success')
            return redirect(url_for('login'))

        return render_template('reset_password.html', token=token)

    except Exception as e:
        flash(f'An error occurred: {str(e)}', 'error')
        return redirect(url_for('forgot_password'))

# Update the forgot_password route's URL generation
@app.route('/forgot-password', methods=['GET', 'POST'])
def forgot_password():
    if request.method == 'POST':
        try:
            # Get email from JSON or form data
            if request.is_json:
                email = request.json.get('email')
            else:
                email = request.form.get('email')

            if not email:
                return jsonify({
                    'success': False,
                    'message': 'Email is required'
                }), 400

            # Check if user exists
            user = mongo.db.users.find_one({"email": email})
            
            if not user:
                return jsonify({
                    'success': False,
                    'message': 'No account found with this email address'
                }), 404

            # Generate token
            token = generate_reset_token(email)

            # Create reset URL - Note the hyphen in 'reset-password'
            reset_url = url_for('reset_password', token=token, _external=True)

            try:
                # Send email with SMTP
                sender_email = "rockstarpraveen44@gmail.com"
                sender_password = "uqmt yfwc khzo gkgj"  # Your app password

                # Create message
                msg = MIMEMultipart()
                msg['From'] = 'Audit IQ AI <{}>'.format(sender_email)
                msg['To'] = email
                msg['Subject'] = "Password Reset Request"

                # HTML Content
                html = f"""
                <html>
                    <body style="font-family: Arial, sans-serif; line-height: 1.6;">
                        <div style="max-width: 600px; margin: 0 auto; padding: 20px;">
                            <h2>Password Reset Request</h2>
                            <p>Hello,</p>
                            <p>We received a request to reset your password. Click the link below to reset it:</p>
                            <p><a href="{reset_url}" style="background: #4361ee; color: white; padding: 10px 20px; text-decoration: none; border-radius: 5px;">Reset Password</a></p>
                            <p>If you didn't request this, please ignore this email.</p>
                            <p>Best regards,<br>Audit IQ AI Team</p>
                        </div>
                    </body>
                </html>
                """
                
                msg.attach(MIMEText(html, 'html'))

                print("Connecting to SMTP server...")
                server = smtplib.SMTP('smtp.gmail.com', 587)
                server.starttls()
                
                print("Logging into SMTP server...")
                server.login(sender_email, sender_password)
                
                print("Sending email...")
                server.send_message(msg)
                server.quit()
                print("Email sent successfully!")

                return jsonify({
                    'success': True,
                    'message': 'Password reset link sent to your email'
                })

            except Exception as email_error:
                print(f"Error sending email: {str(email_error)}")
                return jsonify({
                    'success': False,
                    'message': f'Failed to send email: {str(email_error)}'
                }), 500

        except Exception as e:
            print(f"Error in forgot_password route: {str(e)}")
            return jsonify({
                'success': False,
                'message': f'An error occurred: {str(e)}'
            }), 500

    return render_template('forgot_password.html')


@app.context_processor
def inject_client_id():
    client_id = request.view_args.get('client_id') if 'client_id' in request.view_args else None
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    return dict(client_id=client_id, start_date=start_date, end_date=end_date)

def generate_reset_token(email):
    serializer = URLSafeTimedSerializer(app.config['SECRET_KEY'])
    return serializer.dumps(email, salt='password-reset-salt')

def verify_reset_token(token):
    serializer = URLSafeTimedSerializer(app.config['SECRET_KEY'])
    try:
        email = serializer.loads(token, salt='password-reset-salt', max_age=3600)  # 1 hour expiry
        return email
    except:
        return None

@app.route('/dashboard')
def dashboard():
    if 'logged_in' not in session:
        flash('Please log in to access the dashboard')
        return redirect(url_for('login'))

    # Retrieve clients created by the logged-in user only
    user_id = session.get('user_id')
    clients = list(mongo.db.clients.find({"user_id": user_id}))

    return render_template('dashboard.html', clients=clients)

@app.context_processor
def utility_processor():
    def get_client_id_from_url():
        # Helper function to get client_id from URL
        path_parts = request.path.split('/')
        try:
            client_idx = path_parts.index('client')
            if len(path_parts) > client_idx + 1:
                return path_parts[client_idx + 1]
        except ValueError:
            pass
        return None
    return dict(get_client_id_from_url=get_client_id_from_url)

@app.route('/api/clients/<client_id>', methods=['DELETE'])
def delete_client_api(client_id):
    try:
        # Check if client exists
        client = mongo.db.clients.find_one({"_id": ObjectId(client_id)})
        if not client:
            return jsonify({
                "success": False,
                "message": "Client not found"
            }), 404

        # Delete the client
        result = mongo.db.clients.delete_one({"_id": ObjectId(client_id)})
        
        if result.deleted_count > 0:
            return jsonify({
                "success": True,
                "message": "Client deleted successfully"
            })
        else:
            return jsonify({
                "success": False,
                "message": "Failed to delete client"
            }), 500

    except Exception as e:
        return jsonify({
            "success": False,
            "message": f"An error occurred: {str(e)}"
        }), 500


@app.route('/add_client', methods=['GET', 'POST'])
def add_client():
    if 'logged_in' not in session:
        flash('Please log in to add a client')
        return redirect(url_for('login'))
    
    if request.method == 'POST':
        # Retrieve the user_id from the session
        user_id = session.get('user_id')
        if not user_id:
            flash("User ID not found. Please log in again.", "error")
            return redirect(url_for('login'))
        
        # Collect client data from the form
        client_data = {
            "user_id": user_id,  # Associate the client with the logged-in user's ID
            "company_name": request.form.get('company_name'),
            "business_registration_number": request.form.get('business_registration_number'),
            "gstin": request.form.get('gstin'),
            "pan_tan": request.form.get('pan_tan'),
            "registered_address": request.form.get('registered_address'),
            "operating_addresses": request.form.get('operating_addresses'),
            "key_management_personnel": request.form.get('key_management_personnel'),
            "ownership_structure": request.form.get('ownership_structure'),
            "date_of_incorporation": request.form.get('date_of_incorporation'),
            "authorized_signatories": request.form.get('authorized_signatories'),
            "bank_account_details": request.form.get('bank_account_details'),
            "industry_type": request.form.get('industry_type'),
            "other_industry_type": request.form.get('other_industry_type'),
            "nature_of_products": request.form.get('nature_of_products'),
            "key_raw_materials": request.form.get('key_raw_materials'),
            "key_markets": request.form.get('key_markets'),
            "major_competitors": request.form.get('major_competitors'),
            "manufacturing_processes": request.form.get('manufacturing_processes'),
            "key_technology": request.form.get('key_technology'),
            "capacity_utilization": request.form.get('capacity_utilization'),
            "regulatory_requirements": request.form.get('regulatory_requirements'),
            "product_categories": request.form.get('product_categories'),
            "top_selling_products": request.form.get('top_selling_products'),
            "production_capacity": request.form.get('production_capacity'),
            "lead_time": request.form.get('lead_time'),
            "product_differentiation": request.form.get('product_differentiation'),
            "product_lifecycle": request.form.get('product_lifecycle'),
            "pricing_model": request.form.get('pricing_model'),
            "after_sales_services": request.form.get('after_sales_services'),
            "ownership_structure_details": request.form.get('ownership_structure_details'),
            "board_of_directors": request.form.get('board_of_directors'),
            "corporate_governance": request.form.get('corporate_governance'),
            "legal_advisors_auditors": request.form.get('legal_advisors_auditors'),
            "key_policies": request.form.get('key_policies'),
            "revenue_streams": request.form.get('revenue_streams'),
            "cost_structure": request.form.get('cost_structure'),
            "profit_margins": request.form.get('profit_margins'),
            "capital_structure": request.form.get('capital_structure'),
            "financial_ratios": request.form.get('financial_ratios'),
            "tax_compliance": request.form.get('tax_compliance'),
            "employee_count": request.form.get('employee_count'),
            "employee_turnover_rate": request.form.get('employee_turnover_rate'),
            "compensation_benefits": request.form.get('compensation_benefits'),
            "training_programs": request.form.get('training_programs'),
            "performance_management": request.form.get('performance_management'),
            "workforce_diversity": request.form.get('workforce_diversity'),
            "labor_union": request.form.get('labor_union'),
            "erp_software": request.form.get('erp_software'),
            "data_security_measures": request.form.get('data_security_measures'),
            "backup_disaster_recovery": request.form.get('backup_disaster_recovery'),
            "system_integrations": request.form.get('system_integrations'),
            "user_access_control": request.form.get('user_access_control'),
            "key_risks": request.form.get('key_risks'),
            "regulatory_compliance": request.form.get('regulatory_compliance'),
            "insurance_policies": request.form.get('insurance_policies'),
            "internal_controls": request.form.get('internal_controls'),
            "created_by": session['username']  # Track who created the client
        }
        
        # Insert the new client into the MongoDB clients collection
        mongo.db.clients.insert_one(client_data)
        
        flash('Client added successfully!')
        return redirect(url_for('dashboard'))
    
    return render_template('add_client.html')


# API route for getting clients with optional sorting and search
@app.route('/api/clients', methods=['GET'])
def api_get_clients():
    if 'user_id' not in session:
        return jsonify({"error": "User not logged in"}), 401

    user_id = session['user_id']
    search_query = request.args.get('search', '')
    sort_order = request.args.get('order', 'asc')

    # Build the query to fetch only clients for the logged-in user
    query = {"user_id": user_id}
    
    # Add search filter if provided
    if search_query:
        query["company_name"] = {"$regex": search_query, "$options": "i"}  # Case-insensitive search

    # Fetch clients from MongoDB with applied filters
    clients_cursor = mongo.db.clients.find(query)

    # Sort the clients
    clients = list(clients_cursor)
    clients.sort(key=lambda x: x['company_name'], reverse=(sort_order == 'desc'))

    # Convert ObjectId to string for JSON serialization
    for client in clients:
        client['_id'] = str(client['_id'])

    return jsonify(clients)



@app.route('/api/clients/<client_id>', methods=['GET'])
def api_get_client(client_id):
    client = get_client_by_id(mongo, client_id)
    if client:
        client['_id'] = str(client['_id'])  # Convert ObjectId to string for JSON serialization
    return jsonify(client)

@app.route('/api/clients/<client_id>', methods=['DELETE'])
def api_delete_client(client_id):
    message = delete_client(mongo, client_id)
    return jsonify({"message": message}), 200

@app.route('/add_audit', methods=['GET', 'POST'])
def add_audit():
    if 'logged_in' not in session:
        flash('Please log in to add audit')
        return redirect(url_for('login'))
    if request.method == 'POST':
        audit_data = {
            "name": request.form['audit_name'],
            "details": request.form['audit_details'],
            "created_by": session['username']
        }
        
        mongo.db.audits.insert_one(audit_data)
        flash('Audit added successfully!')
        return redirect(url_for('dashboard'))
    return render_template('add_audit.html')



@app.route('/client/edit/<client_id>', methods=['GET', 'POST'])
def edit_client(client_id):
    # Ensure the user is logged in
    if 'username' not in session:
        flash("Please log in to access this page.")
        return redirect(url_for('login'))

    client = get_client_by_id(mongo, client_id)
    
    if request.method == 'POST':
        # Collect updated data from the form
        updated_data = {
            "company_name": request.form.get('company_name'),
            "business_registration_number": request.form.get('business_registration_number'),
            "gstin": request.form.get('gstin'),
            "pan_tan": request.form.get('pan_tan'),
            "registered_address": request.form.get('registered_address'),
            "operating_addresses": request.form.get('operating_addresses'),
            "key_management_personnel": request.form.get('key_management_personnel'),
            "ownership_structure": request.form.get('ownership_structure'),
            "date_of_incorporation": request.form.get('date_of_incorporation'),
            "authorized_signatories": request.form.get('authorized_signatories'),
            "bank_account_details": request.form.get('bank_account_details'),
            "industry_type": request.form.get('industry_type'),
            "If Other_Specify": request.form.get('other_industry_type'),
            # Products and Services
            # Products and Services
            "nature_of_products": request.form.get('nature_of_products'),
            "key_raw_materials": request.form.get('key_raw_materials'),
            "key_markets": request.form.get('key_markets'),
            "major_competitors": request.form.get('major_competitors'),
            "manufacturing_processes": request.form.get('manufacturing_processes'),
            "key_technology": request.form.get('key_technology'),
            "capacity_utilization": request.form.get('capacity_utilization'),
            "regulatory_requirements": request.form.get('regulatory_requirements'),
            "product_categories": request.form.get('product_categories'),
            "top_selling_products": request.form.get('top_selling_products'),
            "production_capacity": request.form.get('production_capacity'),
            "lead_time": request.form.get('lead_time'),
            "product_differentiation": request.form.get('product_differentiation'),
            "product_lifecycle": request.form.get('product_lifecycle'),
            "pricing_model": request.form.get('pricing_model'),
            "after_sales_services": request.form.get('after_sales_services'),

            # Ownership and Governance Structure
            "ownership_structure_details": request.form.get('ownership_structure_details'),
            "board_of_directors": request.form.get('board_of_directors'),
            "corporate_governance": request.form.get('corporate_governance'),
            "legal_advisors_auditors": request.form.get('legal_advisors_auditors'),
            "key_policies": request.form.get('key_policies'),

            # Financial Overview
            "revenue_streams": request.form.get('revenue_streams'),
            "cost_structure": request.form.get('cost_structure'),
            "profit_margins": request.form.get('profit_margins'),
            "capital_structure": request.form.get('capital_structure'),
            "financial_ratios": request.form.get('financial_ratios'),
            "tax_compliance": request.form.get('tax_compliance'),

            # Human Resources and Payroll
            "employee_count": request.form.get('employee_count'),
            "employee_turnover_rate": request.form.get('employee_turnover_rate'),
            "compensation_benefits": request.form.get('compensation_benefits'),
            "training_programs": request.form.get('training_programs'),
            "performance_management": request.form.get('performance_management'),
            "workforce_diversity": request.form.get('workforce_diversity'),
            "labor_union": request.form.get('labor_union'),

            # IT Systems and Infrastructure
            "erp_software": request.form.get('erp_software'),
            "data_security_measures": request.form.get('data_security_measures'),
            "backup_disaster_recovery": request.form.get('backup_disaster_recovery'),
            "system_integrations": request.form.get('system_integrations'),
            "user_access_control": request.form.get('user_access_control'),

            # Risk Management and Compliance
            "key_risks": request.form.get('key_risks'),
            "regulatory_compliance": request.form.get('regulatory_compliance'),
            "insurance_policies": request.form.get('insurance_policies'),
            "internal_controls": request.form.get('internal_controls'),

            # Meta data
            "created_by": session.get('username', 'default_user')  # Provide default in case 'username' is missing
        }
        
        # Update client in MongoDB
        update_client(mongo, client_id, updated_data)
        
        flash('Client updated successfully!')
        return redirect(url_for('dashboard'))
    
    return render_template('edit_client.html', client=client)


@app.route('/client/<client_id>/open', methods=['GET'])
def open_client_page(client_id):
    client = get_client_by_id(mongo, client_id)
    logging.debug(f"Client data: {client}")
    
    if not client:
        flash("Client not found", "error")
        return redirect(url_for("dashboard"))

    return render_template('client_page.html', client=client)

@app.route('/client/<client_id>/company_information', methods=['GET'])
def company_information(client_id):
    client = get_client_by_id(mongo, client_id)
    return render_template('company_information.html', client=client)


@app.route('/client/<client_id>/audit_findings', methods=['GET'])
def audit_findings(client_id):
    client = get_client_by_id(mongo, client_id)
    return render_template('audit_findings.html', client=client)

@app.route('/client/<client_id>/review_approval', methods=['GET'])
def review_approval(client_id):
    client = get_client_by_id(mongo, client_id)
    return render_template('review_approval.html', client=client)


@app.route('/client/<client_id>/audit_closure', methods=['GET'])
def audit_closure(client_id):
    client = get_client_by_id(mongo, client_id)
    return render_template('audit_closure.html', client=client)

@app.route('/client/<client_id>/manage_users', methods=['GET', 'POST'])
def manage_users(client_id):
    client = mongo.db.clients.find_one({"_id": ObjectId(client_id)})

    if request.method == 'POST':
        user_type = request.form['user_type']  # New field to distinguish user type
        user_data = {
            "username": request.form['username'],
            "role": request.form['role'],
            "email": request.form['email'],
            "phone": request.form['phone'],
            "client_id": client_id  # Associate the user with this client
        }
        
        # Insert data into appropriate collection based on user type
        if user_type == "main":
            mongo.db.users.insert_one(user_data)
            flash('Main user added successfully!')
        elif user_type == "team_member":
            mongo.db.team_users.insert_one(user_data)
            flash('Team member added successfully!')
        else:
            flash('Invalid user type specified.')
        
        return redirect(url_for('manage_users', client_id=client_id))

    # Retrieve main users and team members for this client
    main_users = list(mongo.db.users.find({"client_id": client_id}))
    team_users = list(mongo.db.team_users.find({"client_id": client_id}))

    return render_template('manage_users.html', client=client, main_users=main_users, team_users=team_users)



@app.route('/client/<client_id>/delete_user/<user_id>', methods=['POST'])
def delete_user(client_id, user_id):
    mongo.db.users.delete_one({"_id": ObjectId(user_id)})
    flash('User deleted successfully.')
    return redirect(url_for('manage_users', client_id=client_id))

# View audit planning for a client
@app.route('/client/<client_id>/select_period', methods=['GET', 'POST'])
def select_period(client_id):
    try:
        client = clients_collection.find_one({"_id": ObjectId(client_id)})
        if not client:
            flash("Client not found", "error")
            return redirect(url_for("dashboard"))

        if request.method == 'POST':
            # Check if this is a clone operation or new period creation
            if 'clone_from_period' in request.form:
                # Clone operation
                source_period_id = request.form.get('clone_from_period')
                new_start_date = request.form.get('new_start_date')
                new_end_date = request.form.get('new_end_date')

                # Validate dates
                if not all([source_period_id, new_start_date, new_end_date]):
                    flash('All fields are required for cloning', 'error')
                    return redirect(url_for('select_period', client_id=client_id))

                # Check if target period already exists
                existing_period = client_audit_plan_collection.find_one({
                    "client_id": ObjectId(client_id),
                    "start_date": new_start_date,
                    "end_date": new_end_date
                })

                if existing_period:
                    flash('Audit period already exists for these dates', 'warning')
                else:
                    try:
                        # Get source period data
                        source_period = client_audit_plan_collection.find_one({
                            "_id": ObjectId(source_period_id)
                        })

                        if source_period:
                            # Create new period with cloned data
                            new_period = {
                                "client_id": ObjectId(client_id),
                                "start_date": new_start_date,
                                "end_date": new_end_date,
                                "modified_audit_scope": source_period.get('modified_audit_scope', []),
                                "objectives": source_period.get('objectives', ''),
                                "scope": source_period.get('scope', ''),
                                "cloned_from": source_period_id,
                                "created_at": datetime.now()
                            }
                            client_audit_plan_collection.insert_one(new_period)
                            flash('Audit period cloned successfully!', 'success')
                            return redirect(url_for('audit_planning', 
                                                  client_id=client_id, 
                                                  start_date=new_start_date, 
                                                  end_date=new_end_date))
                        else:
                            flash('Source period not found', 'error')
                    except Exception as e:
                        flash(f'Error during cloning: {str(e)}', 'error')
                        logging.error(f"Clone error: {str(e)}")
            else:
                # New period creation
                start_date = request.form.get('start_date')
                end_date = request.form.get('end_date')

                # Validate dates
                if not start_date or not end_date:
                    flash('Start date and end date are required', 'error')
                    return redirect(url_for('select_period', client_id=client_id))

                try:
                    # Convert string dates to datetime for comparison
                    start_datetime = datetime.strptime(start_date, '%Y-%m-%d')
                    end_datetime = datetime.strptime(end_date, '%Y-%m-%d')

                    if start_datetime > end_datetime:
                        flash('End date must be after start date', 'error')
                        return redirect(url_for('select_period', client_id=client_id))

                    # Check if period already exists
                    existing_period = client_audit_plan_collection.find_one({
                        "client_id": ObjectId(client_id),
                        "start_date": start_date,
                        "end_date": end_date
                    })

                    if existing_period:
                        flash('An audit period already exists for these dates', 'warning')
                    else:
                        # Create new audit period
                        new_period = {
                            "client_id": ObjectId(client_id),
                            "start_date": start_date,
                            "end_date": end_date,
                            "created_at": datetime.now(),
                            "modified_audit_scope": [],  # Initialize empty scope
                            "objectives": "",
                            "scope": "",
                            "status": "New"
                        }
                        
                        client_audit_plan_collection.insert_one(new_period)
                        flash('New audit period created successfully!', 'success')
                        return redirect(url_for('audit_planning', 
                                              client_id=client_id, 
                                              start_date=start_date, 
                                              end_date=end_date))

                except ValueError as ve:
                    flash('Invalid date format', 'error')
                    logging.error(f"Date format error: {str(ve)}")
                except Exception as e:
                    flash(f'Error creating new period: {str(e)}', 'error')
                    logging.error(f"Period creation error: {str(e)}")

        # Fetch all periods for display
        all_periods = list(client_audit_plan_collection.find(
            {"client_id": ObjectId(client_id)},
            sort=[("created_at", -1)]
        ))
        
        # Convert ObjectId to string for template
        client['_id'] = str(client['_id'])
        
        return render_template(
            "select_period.html", 
            client=client, 
            periods=all_periods
        )

    except Exception as e:
        logging.error(f"Error in select_period: {str(e)}")
        flash("An error occurred while processing your request", "error")
        return redirect(url_for("dashboard"))
    
@app.route('/client/<client_id>/delete_period/<period_id>', methods=['DELETE'])
def delete_period(client_id, period_id):
    try:
        # Check if the period exists
        period = client_audit_plan_collection.find_one({
            "_id": ObjectId(period_id),
            "client_id": ObjectId(client_id)
        })

        if not period:
            return jsonify({
                "success": False,
                "error": "Audit period not found"
            }), 404

        # Check if there are any related audit execution records
        execution_records = audit_execution_dynamic.find_one({
            "client_id": ObjectId(client_id),
            "period_start": period['start_date'],
            "period_end": period['end_date']
        })

        if execution_records:
            # Delete related execution records
            audit_execution_dynamic.delete_many({
                "client_id": ObjectId(client_id),
                "period_start": period['start_date'],
                "period_end": period['end_date']
            })

        # Delete the period
        result = client_audit_plan_collection.delete_one({
            "_id": ObjectId(period_id),
            "client_id": ObjectId(client_id)
        })

        if result.deleted_count > 0:
            return jsonify({
                "success": True,
                "message": "Audit period and related data deleted successfully"
            })
        else:
            return jsonify({
                "success": False,
                "error": "Failed to delete audit period"
            }), 500

    except Exception as e:
        logging.error(f"Error deleting audit period: {str(e)}")
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500
    
@app.route('/client/<client_id>/audit_planning')
def audit_planning(client_id):
    try:
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')

        if not start_date or not end_date:
            flash("Please select an audit period first", "warning")
            return redirect(url_for('select_period', client_id=client_id))

        client = clients_collection.find_one({"_id": ObjectId(client_id)})
        if not client:
            flash("Client not found", "error")
            return redirect(url_for("dashboard"))

        # Get or initialize audit plan
        audit_plan = client_audit_plan_collection.find_one({
            "client_id": ObjectId(client_id),
            "start_date": start_date,
            "end_date": end_date
        })

        # Fetch team members from both collections
        main_users = list(mongo.db.users.find({"client_id": client_id}))
        team_users = list(mongo.db.team_users.find({"client_id": client_id}))
        
        # Combine all users into one list
        team_members = main_users + team_users

        # Convert ObjectId to string
        client['_id'] = str(client['_id'])

        return render_template(
            'audit_planning.html',
            client=client,
            audit_plan=audit_plan,
            start_date=start_date,
            end_date=end_date,
            team_members=team_members  # Pass team members to template
        )

    except Exception as e:
        logging.error(f"Error in audit planning: {str(e)}")
        flash("An error occurred while loading audit planning.", "error")
        return redirect(url_for("dashboard"))

@app.route('/client/<client_id>/update_objectives', methods=['POST'])
def update_audit_objectives(client_id):
    start_date = request.form.get('start_date')
    end_date = request.form.get('end_date')

    if not start_date or not end_date:
        flash("Invalid period selected", "error")
        return redirect(url_for('select_period', client_id=client_id))

    # First, verify the audit plan exists for this specific period
    audit_plan = client_audit_plan_collection.find_one({
        "client_id": ObjectId(client_id),
        "start_date": start_date,
        "end_date": end_date
    })

    if not audit_plan:
        flash("No audit plan found for the selected period.", "error")
        return redirect(url_for('select_period', client_id=client_id))

    # Update the specific period's audit plan
    result = client_audit_plan_collection.update_one(
        {
            "client_id": ObjectId(client_id),
            "start_date": start_date,
            "end_date": end_date
        },
        {"$set": {
            "objectives": request.form.get('objectives', ''),
            "scope": request.form.get('scope', ''),
            "modified_audit_scope": request.form.getlist('modified_audit_scope', [])
        }}
    )

    if result.modified_count > 0:
        flash(f"Audit plan for period {start_date} to {end_date} updated successfully!")
    else:
        flash("No changes were made to the audit plan.", "warning")

    return redirect(url_for('audit_planning', 
                          client_id=client_id, 
                          start_date=start_date, 
                          end_date=end_date))

@app.route('/client/<client_id>/update_timeline', methods=['POST'])
def update_audit_timeline(client_id):
    old_start_date = request.form.get('old_start_date')
    old_end_date = request.form.get('old_end_date')
    new_start_date = request.form.get('new_start_date')
    new_end_date = request.form.get('new_end_date')

    # Check if new period already exists
    existing_period = client_audit_plan_collection.find_one({
        "client_id": ObjectId(client_id),
        "start_date": new_start_date,
        "end_date": new_end_date
    })

    if existing_period:
        flash("An audit plan already exists for the new period.", "error")
        return redirect(url_for('audit_planning', 
                              client_id=client_id, 
                              start_date=old_start_date, 
                              end_date=old_end_date))

    # Update the timeline for the specific period
    result = client_audit_plan_collection.update_one(
        {
            "client_id": ObjectId(client_id),
            "start_date": old_start_date,
            "end_date": old_end_date
        },
        {"$set": {
            "start_date": new_start_date,
            "end_date": new_end_date
        }}
    )

    if result.modified_count > 0:
        flash("Timeline updated successfully!")
    else:
        flash("No changes were made to the timeline.", "warning")

    return redirect(url_for('audit_planning', 
                          client_id=client_id, 
                          start_date=new_start_date, 
                          end_date=new_end_date))



# Allocate resources
@app.route('/client/<client_id>/allocate_resources', methods=['POST'], endpoint='allocate_resources')
def allocate_resources(client_id):
    mongo.db.audit_planning.update_one(
        {"client_id": client_id},
        {"$set": {"resources": request.form['resources']}},
        upsert=True
    )
    flash("Resources allocated successfully!")
    return redirect(url_for('audit_planning', client_id=client_id))

# Update risks
@app.route('/client/<client_id>/update_risks', methods=['POST'], endpoint='update_risks')
def update_risks(client_id):
    mongo.db.audit_planning.update_one(
        {"client_id": client_id},
        {"$set": {"risks": request.form['risks']}},
        upsert=True
    )
    flash("Risk assessment updated successfully!")
    return redirect(url_for('audit_planning', client_id=client_id))

# Update audit procedures
@app.route('/client/<client_id>/update_procedures', methods=['POST'], endpoint='update_procedures')
def update_procedures(client_id):
    mongo.db.audit_planning.update_one(
        {"client_id": client_id},
        {"$set": {"procedures": request.form['procedures']}},
        upsert=True
    )
    flash("Audit procedures updated successfully!")
    return redirect(url_for('audit_planning', client_id=client_id))

@app.route('/get_audit_scope/<client_id>')
def get_audit_scope(client_id):
    try:
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')

        logging.debug(f"Fetching audit scope for period {start_date} to {end_date}")

        if not start_date or not end_date:
            return jsonify({
                "success": False,
                "error": "Start date and end date are required"
            }), 400

        # First try to get the period-specific audit plan
        period_plan = client_audit_plan_collection.find_one({
            "client_id": ObjectId(client_id),
            "start_date": start_date,
            "end_date": end_date
        })

        if period_plan and period_plan.get("modified_audit_scope"):
            return jsonify({
                "success": True,
                "scope": period_plan["modified_audit_scope"]
            })

        # If no period-specific plan exists, get the default scope based on industry
        client = clients_collection.find_one({"_id": ObjectId(client_id)})
        if not client:
            return jsonify({"error": "Client not found"}), 404

        industry = client.get("industry") or client.get("industry_type")
        if not industry:
            return jsonify({"error": "Industry type not specified for client"}), 400

        audit_scope = audit_scope_collection.find_one({"industry": industry})
        if not audit_scope:
            return jsonify({"error": "No audit scope data found for this industry"}), 404

        return jsonify({
            "success": True,
            "scope": audit_scope["scope_areas"]
        })

    except Exception as e:
        logging.error(f"Error getting audit scope: {str(e)}")
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/debug/audit_plans/<client_id>')
def debug_audit_plans(client_id):
    """Debug endpoint to view all audit plans for a client"""
    try:
        plans = list(client_audit_plan_collection.find({"client_id": ObjectId(client_id)}))
        return jsonify({
            "success": True,
            "plans": [{
                "start_date": plan["start_date"],
                "end_date": plan["end_date"],
                "scope_count": len(plan.get("modified_audit_scope", [])),
                "modified_audit_scope": plan.get("modified_audit_scope", [])
            } for plan in plans]
        })
    except Exception as e:
        logging.error(f"Error in debug endpoint: {str(e)}")
        return jsonify({"success": False, "error": str(e)}), 500
    
@app.route("/save_audit_scope/<client_id>", methods=["POST"])
def save_audit_scope(client_id):
    try:
        logging.debug("Entered save_audit_scope function")
        
        scope_data = request.json.get("scope", [])
        industry = request.json.get("industry", "")
        
        if not scope_data:
            logging.error("No scope data provided")
            return jsonify({"error": "No scope data provided"}), 400
        if not industry:
            logging.error("Industry type not specified for client")
            return jsonify({"error": "Industry type not specified for client"}), 400

        # Validate and structure each scope item
        validated_scope_data = []
        for item in scope_data:
            validated_scope_data.append({
                "name": item.get("name", ""),
                "details": item.get("details", ""),
                "detailed_procedure": item.get("detailed_procedure", []),
                "status": item.get("status", ""),
                "allocation_team_member": item.get("allocation_team_member", "")
            })
        
        # Logging scope data and industry
        logging.debug(f"Scope Data to Save: {validated_scope_data}")
        logging.debug(f"Industry: {industry}")

        # Save the audit scope and industry type to the client's document
        result = clients_collection.update_one(
            {"_id": ObjectId(client_id)},
            {"$set": {"audit_scope": validated_scope_data, "industry": industry}}
        )
        
        logging.info(f"Audit scope saved successfully for client_id: {client_id}")
        return jsonify({"success": True, "message": "Audit scope saved successfully."})
    except Exception as e:
        logging.error(f"Error in save_audit_scope: {str(e)}")
        return jsonify({"error": str(e)}), 500
    

@app.route('/client/<client_id>/get_modified_audit_plan', methods=['GET'])
def get_modified_audit_plan(client_id):
    try:
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')

        if not start_date or not end_date:
            return jsonify({
                "success": False,
                "error": "Start date and end date are required"
            }), 400

        # Fetch the modified audit plan for the specific period
        audit_plan = client_audit_plan_collection.find_one({
            "client_id": ObjectId(client_id),
            "start_date": start_date,
            "end_date": end_date
        })

        if not audit_plan:
            return jsonify({
                "success": False,
                "error": "No modified audit plan found for this period"
            }), 404

        # Convert ObjectId to string for JSON serialization
        audit_plan['_id'] = str(audit_plan['_id'])
        audit_plan['client_id'] = str(audit_plan['client_id'])

        return jsonify({
            "success": True,
            "scope": audit_plan.get('modified_audit_scope', [])
        })

    except Exception as e:
        logging.error(f"Error getting modified audit plan: {str(e)}")
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500

@app.route('/save_modified_audit_plan/<client_id>', methods=['POST'])
def save_modified_audit_plan(client_id):
    try:
        data = request.json
        start_date = data.get('start_date')
        end_date = data.get('end_date')
        scope_data = data.get('scope', [])
        industry = data.get('industry', '')

        # Validate required fields
        if not all([start_date, end_date, scope_data]):
            return jsonify({
                "success": False,
                "error": "Missing required fields"
            }), 400

        # Convert client_id to ObjectId
        client_id_obj = ObjectId(client_id)

        # Find existing plan
        existing_plan = client_audit_plan_collection.find_one({
            "client_id": client_id_obj,
            "start_date": start_date,
            "end_date": end_date
        })

        plan_data = {
            "client_id": client_id_obj,
            "start_date": start_date,
            "end_date": end_date,
            "modified_audit_scope": scope_data,
            "industry": industry,
            "updated_at": datetime.now()
        }

        if existing_plan:
            # Update existing plan
            result = client_audit_plan_collection.update_one(
                {"_id": existing_plan["_id"]},
                {"$set": plan_data}
            )
        else:
            # Create new plan
            plan_data["created_at"] = datetime.now()
            result = client_audit_plan_collection.insert_one(plan_data)

        # Convert ObjectId to string for response
        if hasattr(result, 'inserted_id'):
            plan_data['_id'] = str(result.inserted_id)
        plan_data['client_id'] = str(plan_data['client_id'])

        return jsonify({
            "success": True,
            "message": "Audit plan saved successfully",
            "plan": json.loads(json_util.dumps(plan_data))
        })

    except Exception as e:
        logging.error(f"Error saving modified audit plan: {str(e)}")
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500

@app.route('/client/<client_id>/get_created_audit_plan')
def get_created_audit_plan(client_id):
    try:
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')

        # Fetch client data
        client = clients_collection.find_one({"_id": ObjectId(client_id)})
        if not client:
            flash("Client not found", "error")
            return redirect(url_for("dashboard"))

        # Get the audit plan for the specific period
        audit_plan = client_audit_plan_collection.find_one({
            "client_id": ObjectId(client_id),
            "start_date": start_date,
            "end_date": end_date
        })

        if not audit_plan and request.headers.get("Accept") == "application/json":
            return jsonify({
                "success": False,
                "error": "No audit plan found for this period"
            }), 404

        # Convert ObjectId to string for JSON serialization
        if client:
            client['_id'] = str(client['_id'])

        return render_template(
            'created_audit_plan.html',
            client=client,
            audit_plan=audit_plan,
            audit_scope=audit_plan.get("modified_audit_scope", []) if audit_plan else [],
            start_date=start_date,
            end_date=end_date
        )

    except Exception as e:
        logging.error(f"Error getting created audit plan: {str(e)}")
        if request.headers.get("Accept") == "application/json":
            return jsonify({"success": False, "error": str(e)}), 500
        flash(f"An error occurred: {str(e)}", "error")
        return redirect(url_for('dashboard'))



@app.route('/client/<client_id>/audit_execution', methods=['GET'])
def audit_execution(client_id):
    try:
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')

        logging.debug(f"Accessing audit_execution with params: client_id={client_id}, start_date={start_date}, end_date={end_date}")

        if not start_date or not end_date:
            flash("Please select an audit period first", "warning")
            return redirect(url_for('select_period', client_id=client_id))

        # Fetch client data
        client = clients_collection.find_one({"_id": ObjectId(client_id)})
        if not client:
            flash("Client not found", "error")
            return redirect(url_for("dashboard"))

        # Get the audit plan for the period
        audit_plan = client_audit_plan_collection.find_one({
            "client_id": ObjectId(client_id),
            "start_date": start_date,
            "end_date": end_date
        })

        if not audit_plan:
            flash("No audit plan found for the selected period", "error")
            return redirect(url_for('select_period', client_id=client_id))

        # Get existing execution data first
        execution_data = list(audit_execution_dynamic.find({
            "client_id": ObjectId(client_id),
            "period_start": start_date,
            "period_end": end_date
        }))

        # Get scope areas from audit plan
        scope_areas = audit_plan.get('modified_audit_scope', [])

        # Create or update execution entries for each scope area
        for scope in scope_areas:
            scope_name = scope.get('name')
            existing_execution = next(
                (exec_data for exec_data in execution_data if exec_data['scope_area'] == scope_name), 
                None
            )

            if not existing_execution:
                new_execution = {
                    "client_id": ObjectId(client_id),
                    "period_start": start_date,
                    "period_end": end_date,
                    "scope_area": scope_name,
                    "tasks": [],
                    "status": scope.get('status', 'Pending'),
                    "allocated_team_member": scope.get('allocation_team_member', ''),
                    "created_at": datetime.now(),
                    "updated_at": datetime.now()
                }
                result = audit_execution_dynamic.insert_one(new_execution)
                new_execution['_id'] = result.inserted_id
                execution_data.append(new_execution)

        # Get team members
        team_users = list(team_users_collection.find({"client_id": client_id}))

        # Get available periods for cloning
        # Get available periods for cloning - modified query
        available_periods = list(audit_execution_dynamic.aggregate([
            {
                "$match": {
                    "client_id": ObjectId(client_id),
                    "$or": [
                        {"period_start": {"$ne": start_date}},
                        {"period_end": {"$ne": end_date}}
                    ]
                }
            },
            {
                "$group": {
                    "_id": {
                        "period_start": "$period_start",
                        "period_end": "$period_end"
                    }
                }
            },
            {
                "$project": {
                    "_id": 0,
                    "period_start": "$_id.period_start",
                    "period_end": "$_id.period_end"
                }
            }
        ]))

        # Convert ObjectId to string
        client['_id'] = str(client['_id'])

        # Convert ObjectIds in execution_data
        for exec_data in execution_data:
            exec_data['_id'] = str(exec_data['_id'])
            exec_data['client_id'] = str(exec_data['client_id'])
            for task in exec_data.get('tasks', []):
                if '_id' in task:
                    task['_id'] = str(task['_id'])

        logging.debug(f"Rendering audit_execution with {len(execution_data)} execution records")

        return render_template(
            'audit_execution.html',
            client=client,
            audit_execution_data=execution_data,
            team_users=team_users,
            start_date=start_date,
            end_date=end_date,
            audit_plan=audit_plan,
            available_periods=available_periods  # Pass the periods to template
        )
    except Exception as e:
        logging.error(f"Error in audit_execution: {str(e)}")
        logging.error(f"Traceback: {traceback.format_exc()}")  # Add detailed error logging
        flash("An error occurred while loading the audit execution page.", "error")
        return redirect(url_for('select_period', client_id=client_id))


@app.route('/client/<client_id>/add_audit_task', methods=['POST'])
def add_audit_task(client_id):
    try:
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')
        
        if not start_date or not end_date:
            return jsonify({"success": False, "error": "Period dates are required"}), 400

        task_data = {
            "_id": ObjectId(),
            "task_name": request.form.get("task_name"),
            "procedure": request.form.get("procedure"),
            "audit_evidence": request.form.get("audit_evidence"),
            "comments": request.form.get("comments"),
            "status": request.form.get("status"),
            "allocated_team_member": request.form.get("allocated_team_member"),
            "created_at": datetime.now(),
            "updated_at": datetime.now()
        }
        
        scope_area = request.form.get("scope_area")

        # Add task to dynamic collection
        result = audit_execution_dynamic.update_one(
            {
                "client_id": ObjectId(client_id),
                "scope_area": scope_area,
                "period_start": start_date,
                "period_end": end_date
            },
            {
                "$push": {"tasks": task_data},
                "$set": {"updated_at": datetime.now()}
            }
        )

        if result.modified_count > 0:
            return jsonify({
                "success": True, 
                "task_id": str(task_data["_id"])
            })
        else:
            return jsonify({
                "success": False, 
                "error": "Failed to add task"
            }), 500

    except Exception as e:
        logging.error(f"Error adding task: {str(e)}")
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/client/<client_id>/update_audit_task/<task_id>', methods=['POST'])
def update_audit_task(client_id, task_id):
    try:
        data = request.json
        logging.info(f"Updating task {task_id} with data: {data}")
        
        # Update the specific task in audit_execution_dynamic
        result = audit_execution_dynamic.update_one(
            {
                "client_id": ObjectId(client_id),
                "tasks._id": ObjectId(task_id)
            },
            {
                "$set": {
                    "tasks.$.task_name": data.get('task_name'),
                    "tasks.$.procedure": data.get('procedure'),
                    "tasks.$.audit_evidence": data.get('audit_evidence'),
                    "tasks.$.comments": data.get('comments'),
                    "tasks.$.status": data.get('status'),
                    "tasks.$.allocated_team_member": data.get('allocated_team_member'),
                    "tasks.$.updated_at": datetime.now()
                }
            }
        )

        if result.modified_count > 0:
            return jsonify({
                "success": True,
                "message": "Task updated successfully"
            })
        else:
            return jsonify({
                "success": False,
                "message": "Task not found or no changes made"
            }), 404

    except Exception as e:
        logging.error(f"Error updating task: {str(e)}")
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/client/<client_id>/delete_audit_task/<task_id>', methods=['DELETE'])
def delete_audit_task(client_id, task_id):
    try:
        logging.info(f"Deleting task {task_id} for client {client_id}")
        
        # First find the document containing the task
        document = audit_execution_dynamic.find_one({
            "client_id": ObjectId(client_id),
            "tasks._id": ObjectId(task_id)
        })
        
        if not document:
            return jsonify({
                "success": False,
                "message": "Task not found"
            }), 404
            
        # Then remove the specific task
        result = audit_execution_dynamic.update_one(
            {
                "client_id": ObjectId(client_id),
                "tasks._id": ObjectId(task_id)
            },
            {
                "$pull": {
                    "tasks": {
                        "_id": ObjectId(task_id)
                    }
                }
            }
        )

        if result.modified_count > 0:
            return jsonify({
                "success": True,
                "message": "Task deleted successfully"
            })
        else:
            return jsonify({
                "success": False,
                "message": "Failed to delete task"
            }), 500

    except Exception as e:
        logging.error(f"Error deleting task: {str(e)}")
        error_details = f"Error deleting task: {str(e)}\nStack trace: {traceback.format_exc()}"
        logging.error(error_details)
        return jsonify({"success": False, "error": str(e)}), 500

from urllib.parse import unquote

@app.route('/client/<client_id>/generate_audit_procedure/<scope_area>', methods=['POST'])
def generate_audit_procedure(client_id, scope_area):
    try:
        # Decode the scope_area to handle special characters like %2F for '/'
        scope_area = unquote(scope_area)
        logging.debug(f"Decoded scope_area: {scope_area}")

        # Fetch the standard procedures for the specified scope_area
        standard_procedures = audit_execution_collection.find_one({"scope_area": scope_area})
        logging.debug(f"Query Result: {standard_procedures}")

        # Handle case where no procedures are found
        if not standard_procedures or not standard_procedures.get('tasks'):
            return jsonify({
                "success": False,
                "error": f"No standard procedures found for scope area: {scope_area}"
            }), 404

        # Process and generate tasks
        standard_tasks = standard_procedures['tasks']
        new_tasks = [
            {
                "_id": ObjectId(),
                "task_name": task.get("task_name", ""),
                "procedure": task.get("procedure", ""),
                "audit_evidence": task.get("audit_evidence", ""),
                "comments": task.get("comments", ""),
                "status": "Pending",
                "allocated_team_member": "",
                "created_at": datetime.now(),
                "updated_at": datetime.now(),
            }
            for task in standard_tasks
        ]

        # Insert the tasks into the audit_execution_dynamic collection
        audit_execution_dynamic.update_one(
            {
                "client_id": ObjectId(client_id),
                "scope_area": scope_area,
                "period_start": request.args.get("start_date"),
                "period_end": request.args.get("end_date")
            },
            {"$set": {"tasks": new_tasks, "updated_at": datetime.now()}},
            upsert=True
        )

        return jsonify({
            "success": True,
            "message": f"Successfully generated {len(new_tasks)} procedures",
            "count": len(new_tasks)
        })

    except Exception as e:
        logging.error(f"Error in generate_audit_procedure: {str(e)}")
        return jsonify({"success": False, "error": str(e)}), 500




# Add this helper function to validate scope area and get collection
def get_audit_collection(scope_area):
    collection_map = {
        "Inventory Management": db.Audit_procedures_inventory_management,
        "Sales and Revenue Recognition": db.Audit_procedures_sales_revenue,
        "Financial Controls and Reporting": db.Audit_procedures_financial_controls,
        "Fixed Asset Management": db.Audit_procedures_fixed_asset_management,
        "Compliance with Laws and Regulations": db.Audit_procedures_compliance,
        "GST Compliance": db.Audit_procedures_gst_compliance,
        "HR Department Compliance and Operations Audit": db.Audit_procedures_hr_compliance,
        "Information Technology Systems and Controls": db.Audit_procedures_it_controls,
        "Risk Management and Internal Controls": db.Audit_procedures_risk_management,
        "Production and Operations": db.Audit_procedures_production_operations,
        "Health, Safety, and Environmental Compliance": db.Audit_procedures_hse_compliance  # Add this line
    }
    return collection_map.get(scope_area)
    
@app.route('/client/<client_id>/audit_reporting')
def audit_reporting(client_id):
    try:
        # Get period dates from query parameters
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')

        logging.debug(f"Accessing audit_reporting with params: client_id={client_id}, start_date={start_date}, end_date={end_date}")

        # Validate that period dates are provided
        if not start_date or not end_date:
            flash("Please select an audit period first", "warning")
            return redirect(url_for('select_period', client_id=client_id))

        # Fetch client data
        client = clients_collection.find_one({"_id": ObjectId(client_id)})
        if not client:
            flash("Client not found", "error")
            return redirect(url_for("dashboard"))

        # Attempt to fetch data from the audit_execution_dynamic collection
        execution_data = list(audit_execution_dynamic.find({
            "client_id": ObjectId(client_id),
            "period_start": start_date,
            "period_end": end_date
        }))

        if not execution_data:
            # If no data is found in audit_execution_dynamic, fetch from client_audit_plan_collection
            audit_plan = client_audit_plan_collection.find_one({
                "client_id": ObjectId(client_id),
                "start_date": start_date,
                "end_date": end_date
            })

            if audit_plan and audit_plan.get("modified_audit_scope"):
                execution_data = []
                for scope in audit_plan.get("modified_audit_scope", []):
                    execution_entry = {
                        "client_id": ObjectId(client_id),
                        "period_start": start_date,
                        "period_end": end_date,
                        "scope_area": scope.get("name", ""),
                        "tasks": [{
                            "_id": ObjectId(),
                            "task_name": task,
                            "procedure": task,
                            "audit_evidence": "",
                            "comments": "",
                            "status": scope.get("status", "Pending"),
                            "allocated_team_member": scope.get("allocation_team_member", ""),
                            "created_at": datetime.now(),
                            "updated_at": datetime.now()
                        } for task in scope.get("detailed_procedure", [])]
                    }
                    execution_data.append(execution_entry)

                # Optionally, persist execution data in audit_execution_dynamic for future use
                audit_execution_dynamic.insert_many(execution_data)
                logging.info("Inserted execution data into audit_execution_dynamic collection.")

        # Calculate statistics for reporting
        stats = {
            "total_tasks": 0,
            "completed_tasks": 0,
            "in_progress_tasks": 0,
            "pending_tasks": 0
        }

        for execution in execution_data:
            for task in execution.get('tasks', []):
                stats["total_tasks"] += 1
                status = task.get('status', 'Pending')
                if status == 'Completed':
                    stats["completed_tasks"] += 1
                elif status == 'In Progress':
                    stats["in_progress_tasks"] += 1
                else:
                    stats["pending_tasks"] += 1

        # Convert ObjectId to string for template compatibility
        client['_id'] = str(client['_id'])

        logging.debug(f"Rendering audit_reporting with {len(execution_data)} execution records and stats: {stats}")

        # Render template with all necessary data
        return render_template(
            'audit_reporting.html',
            client=client,
            audit_execution_data=execution_data,
            start_date=start_date,
            end_date=end_date,
            total_tasks=stats["total_tasks"],
            completed_tasks=stats["completed_tasks"],
            in_progress_tasks=stats["in_progress_tasks"],
            pending_tasks=stats["pending_tasks"]
        )

    except Exception as e:
        # Log the error and redirect to dashboard with an error message
        logging.error(f"Error in audit_reporting: {str(e)}")
        import traceback
        logging.error(traceback.format_exc())
        flash("An error occurred while loading the audit report.", "error")
        return redirect(url_for("dashboard"))


@app.route('/client/<client_id>/generate_custom_report', methods=['POST'])
def generate_custom_report(client_id):
    try:
        data = request.json
        format_type = data.get('format')
        start_date = data.get('startDate')
        end_date = data.get('endDate')

        # Fetch client info
        client = clients_collection.find_one({"_id": ObjectId(client_id)})
        if not client:
            return jsonify({"error": "Client not found"}), 404

        # Fetch execution data
        execution_data = list(audit_execution_dynamic.find({
            "client_id": ObjectId(client_id),
            "period_start": start_date,
            "period_end": end_date
        }))

        if format_type == 'excel':
            return generate_excel_report(client, execution_data, start_date, end_date)
        else:
            return generate_detailed_word_report(client, execution_data, start_date, end_date)

    except Exception as e:
        logging.error(f"Error generating custom report: {str(e)}")
        return jsonify({"error": str(e)}), 500

def generate_detailed_word_report(client, execution_data, start_date, end_date):
    try:
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'Audit Report - {client["company_name"]}', level=1)
        title.alignment = 1  # Center alignment
        
        # Add period information
        doc.add_paragraph(f'Audit Period: {start_date} to {end_date}')
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=2)
        
        # Calculate statistics
        total_tasks = completed_tasks = in_progress_tasks = pending_tasks = 0
        for execution in execution_data:
            for task in execution.get('tasks', []):
                total_tasks += 1
                status = task.get('status', 'Pending')
                if status == 'Completed':
                    completed_tasks += 1
                elif status == 'In Progress':
                    in_progress_tasks += 1
                else:
                    pending_tasks += 1

        # Add summary table
        summary_table = doc.add_table(rows=5, cols=2)
        summary_table.style = 'Table Grid'
        
        summary_data = [
            ('Total Tasks', str(total_tasks)),
            ('Completed Tasks', str(completed_tasks)),
            ('In Progress Tasks', str(in_progress_tasks)),
            ('Pending Tasks', str(pending_tasks)),
            ('Completion Rate', f"{(completed_tasks/total_tasks*100 if total_tasks > 0 else 0):.1f}%")
        ]
        
        for i, (label, value) in enumerate(summary_data):
            row = summary_table.rows[i].cells
            row[0].text = label
            row[1].text = value
            row[0].paragraphs[0].runs[0].font.bold = True

        doc.add_paragraph()  # Add spacing
        
        # Detailed Findings
        doc.add_heading('Detailed Findings', level=2)
        
        for execution in execution_data:
            scope_area = execution.get('scope_area', '')
            tasks = execution.get('tasks', [])
            
            if tasks:
                doc.add_heading(scope_area, level=3)
                
                # Create table for tasks
                table = doc.add_table(rows=1, cols=6)
                table.style = 'Table Grid'
                
                # Set header
                header_cells = table.rows[0].cells
                headers = ['Task', 'Procedure', 'Audit Evidence', 'Comments', 'Status', 'Team Member']
                for i, header in enumerate(headers):
                    header_cells[i].text = header
                    header_cells[i].paragraphs[0].runs[0].font.bold = True
                
                # Add tasks
                for task in tasks:
                    row_cells = table.add_row().cells
                    row_cells[0].text = task.get('task_name', '')
                    row_cells[1].text = task.get('procedure', '')
                    row_cells[2].text = task.get('audit_evidence', '')
                    row_cells[3].text = task.get('comments', '')
                    row_cells[4].text = task.get('status', '')
                    row_cells[5].text = task.get('allocated_team_member', '')

                doc.add_paragraph()  # Add spacing between tables
        
        # Recommendations
        doc.add_heading('Recommendations', level=2)
        recommendations = doc.add_paragraph()
        recommendations.add_run("Based on the audit findings, the following recommendations are provided:").bold = True
        
        for execution in execution_data:
            for task in execution.get('tasks', []):
                if task.get('status') != 'Completed':
                    doc.add_paragraph(
                        f"• {task.get('task_name', '')}: Complete the pending task and ensure proper documentation.",
                        style='List Bullet'
                    )

        # Save to BytesIO
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)

        return send_file(
            doc_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'Audit_Report_{client["company_name"]}_{start_date}_to_{end_date}.docx'
        )

    except Exception as e:
        logging.error(f"Error generating Word report: {str(e)}")
        raise


def generate_excel_report(client, execution_data, start_date, end_date):
    try:
        wb = Workbook()
        
        # Create Summary Sheet
        ws_summary = wb.active
        ws_summary.title = "Executive Summary"
        
        # Add title and period info
        ws_summary['A1'] = f'Audit Report - {client["company_name"]}'
        ws_summary['A2'] = f'Period: {start_date} to {end_date}'
        ws_summary['A3'] = f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}'
        
        # Calculate statistics
        total_tasks = completed_tasks = in_progress_tasks = pending_tasks = 0
        for execution in execution_data:
            for task in execution.get('tasks', []):
                total_tasks += 1
                status = task.get('status', 'Pending')
                if status == 'Completed':
                    completed_tasks += 1
                elif status == 'In Progress':
                    in_progress_tasks += 1
                else:
                    pending_tasks += 1

        # Add statistics
        stats_data = [
            ['Statistics', 'Count'],
            ['Total Tasks', total_tasks],
            ['Completed Tasks', completed_tasks],
            ['In Progress Tasks', in_progress_tasks],
            ['Pending Tasks', pending_tasks],
            ['Completion Rate', f"{(completed_tasks/total_tasks*100 if total_tasks > 0 else 0):.1f}%"]
        ]
        
        for row in stats_data:
            ws_summary.append(row)
        
        # Create Detailed Findings Sheet
        ws_details = wb.create_sheet("Detailed Findings")
        headers = ['Scope Area', 'Task Name', 'Procedure', 'Audit Evidence', 
                  'Comments', 'Status', 'Team Member']
        ws_details.append(headers)
        
        # Add task details
        for execution in execution_data:
            scope_area = execution.get('scope_area', '')
            for task in execution.get('tasks', []):
                row_data = [
                    scope_area,
                    task.get('task_name', ''),
                    task.get('procedure', ''),
                    task.get('audit_evidence', ''),
                    task.get('comments', ''),
                    task.get('status', ''),
                    task.get('allocated_team_member', '')
                ]
                ws_details.append(row_data)
        
        # Create Recommendations Sheet
        ws_recommendations = wb.create_sheet("Recommendations")
        ws_recommendations.append(['Task Name', 'Status', 'Recommendation'])
        
        for execution in execution_data:
            for task in execution.get('tasks', []):
                if task.get('status') != 'Completed':
                    ws_recommendations.append([
                        task.get('task_name', ''),
                        task.get('status', ''),
                        'Complete the pending task and ensure proper documentation.'
                    ])
        
        # Save to BytesIO
        excel_buffer = io.BytesIO()
        wb.save(excel_buffer)
        excel_buffer.seek(0)

        return send_file(
            excel_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'Audit_Report_{client["company_name"]}_{start_date}_to_{end_date}.xlsx'
        )

    except Exception as e:
        logging.error(f"Error generating Excel report: {str(e)}")
        raise





@app.route('/client/<client_id>/created_audit_plan')
def created_audit_plan(client_id):
    try:
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')

        if not start_date or not end_date:
            flash("Please select an audit period first", "warning")
            return redirect(url_for('select_period', client_id=client_id))

        client = clients_collection.find_one({"_id": ObjectId(client_id)})
        if not client:
            flash("Client not found", "error")
            return redirect(url_for("dashboard"))

        # Get the audit plan
        audit_plan = client_audit_plan_collection.find_one({
            "client_id": ObjectId(client_id),
            "start_date": start_date,
            "end_date": end_date
        })

        # Convert ObjectId to string
        client['_id'] = str(client['_id'])

        return render_template(
            'created_audit_plan.html',
            client=client,
            audit_plan=audit_plan,
            audit_scope=audit_plan.get("modified_audit_scope", []) if audit_plan else [],
            start_date=start_date,
            end_date=end_date
        )

    except Exception as e:
        logging.error(f"Error loading created audit plan: {str(e)}")
        flash("An error occurred while loading the audit plan.", "error")
        return redirect(url_for("dashboard"))

def get_storage_info():
    try:
        # Get all files from GridFS
        files = db['fs.files'].find()
        
        # Calculate total storage used
        total_size = sum(file.get('length', 0) for file in files)
        
        # Count total number of files
        file_count = db['fs.files'].count_documents({})
        
        # Convert bytes to more readable format
        def convert_size(size_bytes):
            for unit in ['B', 'KB', 'MB', 'GB']:
                if size_bytes < 1024.0:
                    return f"{size_bytes:.2f} {unit}"
                size_bytes /= 1024.0
            return f"{size_bytes:.2f} TB"

        return {
            'total_size': convert_size(total_size),
            'file_count': file_count,
            'raw_size': total_size
        }
    except Exception as e:
        logging.error(f"Error calculating storage info: {str(e)}")
        return {
            'total_size': '0 MB',
            'file_count': 0,
            'raw_size': 0
        }

@app.route('/client/<client_id>/file_management')
def file_management(client_id):
    if 'logged_in' not in session:
        flash('Please log in to access file management', 'error')
        return redirect(url_for('login'))

    user_id = session.get('user_id')
    if not user_id:
        flash('User session error', 'error')
        return redirect(url_for('login'))

    try:
        # Fetch files for this specific client
        files = db['fs.files'].find({"metadata.client_id": client_id})
        client = clients_collection.find_one({"_id": ObjectId(client_id)})
        
        # Get storage info for all user's clients
        storage_info = get_user_storage_info(user_id)
        
        return render_template('index.html', 
                             files=files, 
                             client=client, 
                             storage_info=storage_info,
                             is_user_view=True)
    except Exception as e:
        logging.error(f"Error in file management: {str(e)}")
        flash('Error accessing file management', 'error')
        return redirect(url_for('dashboard'))

# Add this constant at the top of your app.py
MAX_FILE_SIZE = 15 * 1024 * 1024   # 15MB in bytes
STORAGE_LIMIT = 50 * 1024 * 1024  # 50MB in bytes

def check_storage_limit(new_file_size):
    """Check if adding new file would exceed storage limit of 100MB"""
    try:
        # Calculate current total storage used
        total_size = 0
        for grid_out in fs.find():
            total_size += grid_out.length
            
        # Convert to MB for readable message
        current_usage_mb = total_size/(1024*1024)
            
        # Check if adding new file would exceed limit
        if (total_size + new_file_size) > STORAGE_LIMIT:
            return False, f"Storage limit of 50MB exceeded. Current usage: {current_usage_mb:.2f}MB"
        return True, None
        
    except Exception as e:
        logging.error(f"Error checking storage limit: {str(e)}")
        return False, "Error checking storage capacity"
        
def get_user_storage_info(user_id):
    """Calculate storage usage for all clients belonging to a user"""
    try:
        # First get all clients belonging to this user
        user_clients = list(clients_collection.find({"user_id": user_id}))
        client_ids = [str(client['_id']) for client in user_clients]
        
        # Find all files for all clients belonging to this user
        files = db['fs.files'].find({"metadata.client_id": {"$in": client_ids}})
        
        # Calculate total storage used
        total_size = sum(file.get('length', 0) for file in files)
        
        # Count total number of files
        file_count = db['fs.files'].count_documents({"metadata.client_id": {"$in": client_ids}})
        
        # Convert bytes to readable format
        def convert_size(size_bytes):
            for unit in ['B', 'KB', 'MB', 'GB']:
                if size_bytes < 1024.0:
                    return f"{size_bytes:.2f} {unit}"
                size_bytes /= 1024.0
            return f"{size_bytes:.2f} TB"

        return {
            'total_size': convert_size(total_size),
            'file_count': file_count,
            'raw_size': total_size,
            'client_count': len(user_clients)
        }
    except Exception as e:
        logging.error(f"Error calculating user storage info: {str(e)}")
        return {
            'total_size': '0 MB',
            'file_count': 0,
            'raw_size': 0,
            'client_count': 0
        }

def check_user_storage_limit(user_id, new_file_size):
    """Check if adding new file would exceed storage limit of 50MB for the user"""
    try:
        # Get current storage info for this user
        user_storage = get_user_storage_info(user_id)
        current_size = user_storage['raw_size']
            
        # Convert to MB for readable message
        current_usage_mb = current_size/(1024*1024)
            
        # Check if adding new file would exceed limit
        if (current_size + new_file_size) > STORAGE_LIMIT:
            return False, f"Your storage limit of 50MB exceeded. Current usage: {current_usage_mb:.2f}MB across {user_storage['client_count']} clients"
        return True, None
        
    except Exception as e:
        logging.error(f"Error checking user storage limit: {str(e)}")
        return False, "Error checking storage capacity"

@app.route('/client/<client_id>/upload', methods=['POST'])
def upload_file(client_id):
    try:
        if 'logged_in' not in session:
            return jsonify({"success": False, "error": "Please log in"}), 401

        user_id = session.get('user_id')
        if not user_id:
            return jsonify({"success": False, "error": "User ID not found"}), 401

        if 'file' not in request.files:
            return jsonify({"success": False, "error": "No file part"}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({"success": False, "error": "No selected file"}), 400
            
        # Check file size
        file_content = file.read()
        file_size = len(file_content)
        
        # First check individual file size limit
        if file_size > MAX_FILE_SIZE:
            return jsonify({
                "success": False,
                "error": "File size exceeds 15MB limit",
                "size": file_size,
                "max_size": MAX_FILE_SIZE
            }), 413
            
        # Then check user's total storage limit
        can_upload, error_message = check_user_storage_limit(user_id, file_size)
        if not can_upload:
            return jsonify({
                "success": False,
                "error": error_message
            }), 413
            
        # Reset file pointer to beginning after reading
        file.seek(0)
        
        if file:
            try:
                # Add both user_id and client_id to metadata
                metadata = {
                    "user_id": user_id,
                    "client_id": client_id,
                    "upload_date": datetime.now()
                }
                
                file_id = fs.put(file, filename=file.filename, metadata=metadata)
                return jsonify({
                    "success": True, 
                    "message": "File uploaded successfully",
                    "file_id": str(file_id)
                }), 200
            except Exception as e:
                logging.error(f"Error uploading file: {str(e)}")
                return jsonify({
                    "success": False, 
                    "error": "Error uploading file"
                }), 500

    except Exception as e:
        logging.error(f"Upload error: {str(e)}")
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500

    except Exception as e:
        logging.error(f"Upload error: {str(e)}")
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500
@app.route('/client/<client_id>/download/<file_id>', methods=['GET'])
def download_file(client_id, file_id):
    try:
        file_data = fs.get(ObjectId(file_id))
        return file_data.read(), 200, {
            'Content-Type': 'application/octet-stream',
            'Content-Disposition': f'attachment; filename={file_data.filename}'
        }
    except gridfs.errors.NoFile:
        return "File not found", 404

@app.route('/client/<client_id>/delete/<file_id>', methods=['POST'])
def delete_file(client_id, file_id):
    try:
        fs.delete(ObjectId(file_id))
        return redirect(url_for('file_management', client_id=client_id))
    except gridfs.errors.NoFile:
        return {"error": "File not found"}, 404

@app.route('/client/<client_id>/view/<file_id>', methods=['GET'])
def view_file(client_id, file_id):
    try:
        file_data = fs.get(ObjectId(file_id))
        return file_data.read(), 200, {
            'Content-Type': 'application/pdf',  # Assuming PDF files for view
            'Content-Disposition': 'inline; filename=' + file_data.filename
        }
    except gridfs.errors.NoFile:
        return "File not found", 404
    
@app.errorhandler(404)
def page_not_found(e):
    logging.error(f"404 Error: {request.url}")
    logging.error(f"Request Args: {request.args}")
    return render_template('404.html'), 404

@app.before_request
def log_request_info():
    logging.debug('Headers: %s', request.headers)
    logging.debug('Body: %s', request.get_data())


def get_current_period():
    """Helper function to get current audit period from request args"""
    return {
        'start_date': request.args.get('start_date'),
        'end_date': request.args.get('end_date')
    }

# Add to context processor
@app.context_processor
def utility_processor():
    return {
        'get_current_period': get_current_period
    }
@app.errorhandler(404)
def page_not_found(e):
    logging.error(f"404 Error: {request.url}")
    
    try:
        # Try to get client_id from URL
        path_parts = request.path.split('/')
        client_id = None
        try:
            client_idx = path_parts.index('client')
            if len(path_parts) > client_idx + 1:
                client_id = path_parts[client_idx + 1]
        except (ValueError, IndexError):
            pass

        # If we have a client_id, try to get client details
        client = None
        if client_id:
            try:
                client = clients_collection.find_one({"_id": ObjectId(client_id)})
                if client:
                    client['_id'] = str(client['_id'])
            except:
                pass

        # First try to render with client_page template
        try:
            return render_template(
                '404.html',
                client=client,
                error_message="The requested page could not be found.",
                start_date=request.args.get('start_date'),
                end_date=request.args.get('end_date')
            ), 404
        except:
            # If that fails, fall back to error_base template
            return render_template(
                'error_base.html',
                error_code=404,
                error_message="The requested page could not be found."
            ), 404

    except Exception as error:
        logging.error(f"Error handling 404: {str(error)}")
        return "Page not found", 404

@app.errorhandler(500)
def internal_error(e):
    logging.error(f"500 Error: {str(e)}")
    
    try:
        return render_template('500.html'), 500
    except:
        return "Internal server error", 500
    
def get_audit_reporting_url(client_id, start_date=None, end_date=None):
    base_url = url_for('audit_reporting', client_id=client_id)
    if start_date and end_date:
        return f"{base_url}?start_date={start_date}&end_date={end_date}"
    return base_url


# In your app.py, after creating the Flask app
def register_routes():
    app.add_url_rule('/client/<client_id>/audit_reporting', 
                     'audit_reporting', 
                     audit_reporting, 
                     methods=['GET'])

# Call this after creating your Flask app
register_routes()

@app.errorhandler(Exception)
def handle_exception(e):
    logging.error(f"Unhandled Exception: {str(e)}")
    
    if isinstance(e, HTTPException):
        return render_template(
            'error_base.html',
            error_code=e.code,
            error_message=e.description
        ), e.code
    
    return render_template(
        'error_base.html',
        error_code=500,
        error_message="An unexpected error occurred."
    ), 500

@app.route('/debug/audit_data/<client_id>')
def debug_audit_data(client_id):
    try:
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')
        
        execution_data = list(audit_execution_dynamic.find({
            "client_id": ObjectId(client_id),
            "period_start": start_date,
            "period_end": end_date
        }))
        
        audit_plan = client_audit_plan_collection.find_one({
            "client_id": ObjectId(client_id),
            "start_date": start_date,
            "end_date": end_date
        })
        
        return jsonify({
            "execution_data_count": len(execution_data),
            "has_audit_plan": bool(audit_plan),
            "scope_count": len(audit_plan.get("modified_audit_scope", [])) if audit_plan else 0
        })
        
    except Exception as e:
        return jsonify({"error": str(e)})

# Add this template filter for formatting dates
@app.template_filter('format_date')
def format_date(value, format='%Y-%m-%d'):
    if isinstance(value, datetime):
        return value.strftime(format)
    return value

# Add this template filter for status colors
@app.template_filter('status_color')
def status_color(status):
    colors = {
        'Completed': 'success',
        'In Progress': 'warning',
        'Pending': 'secondary',
        'Planned': 'info',
        'Delayed': 'danger'
    }
    return colors.get(status.lower(), 'secondary')

# Add template filter for relative time
@app.template_filter('timeago')
def timeago(value):
    if not isinstance(value, datetime):
        return value
        
    now = datetime.now()
    diff = now - value

    if diff.days > 365:
        years = diff.days // 365
        return f"{years} year{'s' if years != 1 else ''} ago"
    elif diff.days > 30:
        months = diff.days // 30
        return f"{months} month{'s' if months != 1 else ''} ago"
    elif diff.days > 0:
        return f"{diff.days} day{'s' if diff.days != 1 else ''} ago"
    elif diff.seconds > 3600:
        hours = diff.seconds // 3600
        return f"{hours} hour{'s' if hours != 1 else ''} ago"
    elif diff.seconds > 60:
        minutes = diff.seconds // 60
        return f"{minutes} minute{'s' if minutes != 1 else ''} ago"
    else:
        return "just now"



@app.route('/client/<client_id>/schedule-audit', methods=['POST'])
def schedule_audit(client_id):
    try:
        audit_data = {
            "client_id": ObjectId(client_id),
            "area": request.form['auditArea'],
            "start_date": request.form['startDate'],
            "end_date": request.form['endDate'],
            "risk_level": request.form['riskLevel'],
            "team_members": request.form.getlist('teamMembers'),
            "status": "Planned",
            "created_at": datetime.now(),
            "created_by": session.get('username')
        }
        
        client_audit_plan_collection.insert_one(audit_data)
        return jsonify({"success": True})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
    
    
@app.route('/client/<client_id>/export_audit_report', methods=['GET'])
def export_audit_report(client_id):
    try:
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')
        
        # Fetch client and audit data
        client = clients_collection.find_one({"_id": ObjectId(client_id)})
        audit_plan = client_audit_plan_collection.find_one({
            "client_id": ObjectId(client_id),
            "start_date": start_date,
            "end_date": end_date
        })
        
        if not client or not audit_plan:
            return jsonify({"error": "Data not found"}), 404

        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'Audit Plan Report - {client["company_name"]}', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add period information
        doc.add_paragraph(f'Audit Period: {start_date} to {end_date}')
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        
        # Add audit scope table
        doc.add_heading('Audit Scope and Procedures', level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Add header row
        header_cells = table.rows[0].cells
        headers = ['Scope Area', 'Details', 'Detailed Procedures', 'Status', 'Allocated To']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Add data rows
        for scope in audit_plan.get('modified_audit_scope', []):
            row_cells = table.add_row().cells
            row_cells[0].text = scope.get('name', '')
            row_cells[1].text = scope.get('details', '')
            
            # Format procedures as bullet points
            procedures = scope.get('detailed_procedure', [])
            procedures_text = '\n'.join(f'• {p}' for p in procedures)
            row_cells[2].text = procedures_text
            
            row_cells[3].text = scope.get('status', '')
            row_cells[4].text = scope.get('allocation_team_member', '')

        # Save to BytesIO
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)

        return send_file(
            doc_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'Audit_Plan_{client["company_name"]}_{start_date}_to_{end_date}.docx'
        )

    except Exception as e:
        logging.error(f"Error exporting Word report: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/client/<client_id>/export_audit_report_excel', methods=['GET'])
def export_audit_report_excel(client_id):
    try:
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')
        
        # Fetch client and audit data
        client = clients_collection.find_one({"_id": ObjectId(client_id)})
        audit_plan = client_audit_plan_collection.find_one({
            "client_id": ObjectId(client_id),
            "start_date": start_date,
            "end_date": end_date
        })
        
        if not client or not audit_plan:
            return jsonify({"error": "Data not found"}), 404

        # Create Excel workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Audit Plan"
        
        # Add title and period info
        ws['A1'] = f'Audit Plan Report - {client["company_name"]}'
        ws['A2'] = f'Period: {start_date} to {end_date}'
        ws['A3'] = f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}'
        
        # Style the header
        for cell in [ws['A1'], ws['A2'], ws['A3']]:
            cell.font = Font(bold=True)
        
        # Add headers
        headers = ['Scope Area', 'Details', 'Detailed Procedures', 'Status', 'Allocated To']
        ws.append([''] * len(headers))  # Empty row for spacing
        ws.append(headers)
        
        # Style the headers
        header_row = ws[5]  # Accounting for title rows and spacing
        for cell in header_row:
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="E0E0E0", end_color="E0E0E0", fill_type="solid")
        
        # Add data
        for scope in audit_plan.get('modified_audit_scope', []):
            procedures = '\n'.join(scope.get('detailed_procedure', []))
            ws.append([
                scope.get('name', ''),
                scope.get('details', ''),
                procedures,
                scope.get('status', ''),
                scope.get('allocation_team_member', '')
            ])
        
        # Adjust column widths
        for column in ws.columns:
            max_length = 0
            column = list(column)
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = (max_length + 2)
            ws.column_dimensions[column[0].column_letter].width = min(adjusted_width, 50)

        # Save to BytesIO
        excel_buffer = io.BytesIO()
        wb.save(excel_buffer)
        excel_buffer.seek(0)

        return send_file(
            excel_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'Audit_Plan_{client["company_name"]}_{start_date}_to_{end_date}.xlsx'
        )

    except Exception as e:
        logging.error(f"Error exporting Excel report: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/client/<client_id>/clone_audit_execution', methods=['POST'])
def clone_audit_execution(client_id):
    try:
        data = request.json
        source_start_date = data.get('source_start_date')
        source_end_date = data.get('source_end_date')
        new_start_date = data.get('new_start_date')
        new_end_date = data.get('new_end_date')
        overwrite = data.get('overwrite', False)  # New parameter for overwriting

        if not all([source_start_date, source_end_date, new_start_date, new_end_date]):
            return jsonify({
                "success": False,
                "error": "All date fields are required"
            }), 400

        # Check if target period already exists
        existing_execution = audit_execution_dynamic.find_one({
            "client_id": ObjectId(client_id),
            "period_start": new_start_date,
            "period_end": new_end_date
        })

        if existing_execution and not overwrite:
            return jsonify({
                "success": False,
                "error": "Audit execution already exists for the target period",
                "requires_overwrite": True
            }), 409  # 409 Conflict

        # Get source execution data
        source_executions = list(audit_execution_dynamic.find({
            "client_id": ObjectId(client_id),
            "period_start": source_start_date,
            "period_end": source_end_date
        }))

        if not source_executions:
            return jsonify({
                "success": False,
                "error": "No source execution data found"
            }), 404

        # If overwriting, delete existing execution data
        if existing_execution:
            audit_execution_dynamic.delete_many({
                "client_id": ObjectId(client_id),
                "period_start": new_start_date,
                "period_end": new_end_date
            })

        # Clone each scope area's execution data
        new_executions = []
        for execution in source_executions:
            new_execution = {
                "client_id": ObjectId(client_id),
                "period_start": new_start_date,
                "period_end": new_end_date,
                "scope_area": execution['scope_area'],
                "tasks": [
                    {
                        "_id": ObjectId(),
                        "task_name": task['task_name'],
                        "procedure": task['procedure'],
                        "audit_evidence": "",  # Reset evidence for new period
                        "comments": "",  # Reset comments for new period
                        "status": "Pending",  # Reset status for new period
                        "allocated_team_member": task.get('allocated_team_member', ''),
                        "created_at": datetime.now(),
                        "updated_at": datetime.now()
                    }
                    for task in execution.get('tasks', [])
                ],
                "created_at": datetime.now(),
                "updated_at": datetime.now(),
                "cloned_from": str(execution['_id'])
            }
            new_executions.append(new_execution)

        if new_executions:
            audit_execution_dynamic.insert_many(new_executions)
            return jsonify({
                "success": True,
                "message": f"Successfully cloned {len(new_executions)} scope areas with tasks"
            })
        else:
            return jsonify({
                "success": False,
                "error": "No data to clone"
            }), 400

    except Exception as e:
        logging.error(f"Error cloning audit execution: {str(e)}")
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500
    

@app.route('/client/<client_id>/audit_dashboard')
def audit_dashboard(client_id):
    try:
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')

        client = clients_collection.find_one({"_id": ObjectId(client_id)})
        if not client:
            flash("Client not found", "error")
            return redirect(url_for("dashboard"))

        audit_periods = list(client_audit_plan_collection.find(
            {"client_id": ObjectId(client_id)},
            sort=[("start_date", -1)]
        ))

        # Filter execution data by period if selected
        execution_query = {"client_id": ObjectId(client_id)}
        if start_date and end_date:
            execution_query.update({
                "period_start": start_date,
                "period_end": end_date
            })
            execution_data = list(audit_execution_dynamic.find(execution_query))
        else:
            # If no period selected, get the most recent period's data
            latest_period = audit_periods[0] if audit_periods else None
            if latest_period:
                execution_data = list(audit_execution_dynamic.find({
                    "client_id": ObjectId(client_id),
                    "period_start": latest_period['start_date'],
                    "period_end": latest_period['end_date']
                }))
                start_date = latest_period['start_date']
                end_date = latest_period['end_date']
            else:
                execution_data = []

        # Calculate statistics for filtered data
        total_tasks = 0
        completed_tasks = 0
        in_progress_tasks = 0
        pending_tasks = 0
        scope_completion = {}

        for execution in execution_data:
            scope_name = execution.get('scope_area')
            scope_tasks = execution.get('tasks', [])
            
            if scope_name not in scope_completion:
                scope_completion[scope_name] = {
                    'total': 0, 'completed': 0, 
                    'in_progress': 0, 'pending': 0
                }

            for task in scope_tasks:
                total_tasks += 1
                scope_completion[scope_name]['total'] += 1
                
                status = task.get('status', 'Pending')
                if status == 'Completed':
                    completed_tasks += 1
                    scope_completion[scope_name]['completed'] += 1
                elif status == 'In Progress':
                    in_progress_tasks += 1
                    scope_completion[scope_name]['in_progress'] += 1
                else:
                    pending_tasks += 1
                    scope_completion[scope_name]['pending'] += 1

        # Calculate scope percentages
        scope_percentages = {
            scope: round((stats['completed'] / stats['total'] * 100), 2) 
            if stats['total'] > 0 else 0
            for scope, stats in scope_completion.items()
        }

        return render_template(
            'audit_dashboard.html',
            client=client,
            audit_periods=audit_periods,
            start_date=start_date,
            end_date=end_date,
            total_tasks=total_tasks,
            completed_tasks=completed_tasks,
            in_progress_tasks=in_progress_tasks,
            pending_tasks=pending_tasks,
            scope_completion=scope_completion,
            scope_percentages=scope_percentages,
            execution_data=execution_data
        )

    except Exception as e:
        logging.error(f"Error in audit dashboard: {str(e)}")
        flash("An error occurred while loading the dashboard.", "error")
        return redirect(url_for("dashboard"))

    
@app.route('/client/<client_id>/export_dashboard_pdf', methods=['POST'])
def export_dashboard_pdf(client_id):
   try:
       start_date = request.args.get('start_date')
       end_date = request.args.get('end_date')

       # Fetch client data
       client = clients_collection.find_one({"_id": ObjectId(client_id)})
       if not client:
           return jsonify({"error": "Client not found"}), 404

       # Get execution data for the period
       execution_data = list(audit_execution_dynamic.find({
           "client_id": ObjectId(client_id),
           "period_start": start_date,
           "period_end": end_date
       }))

       # Calculate statistics
       total_tasks = 0
       completed_tasks = 0
       in_progress_tasks = 0
       pending_tasks = 0
       scope_completion = {}

       for execution in execution_data:
           scope_name = execution.get('scope_area')
           scope_tasks = execution.get('tasks', [])
           
           if scope_name not in scope_completion:
               scope_completion[scope_name] = {
                   'total': 0, 'completed': 0, 
                   'in_progress': 0, 'pending': 0
               }

           for task in scope_tasks:
               total_tasks += 1
               scope_completion[scope_name]['total'] += 1
               
               status = task.get('status', 'Pending')
               if status == 'Completed':
                   completed_tasks += 1
                   scope_completion[scope_name]['completed'] += 1
               elif status == 'In Progress':
                   in_progress_tasks += 1
                   scope_completion[scope_name]['in_progress'] += 1
               else:
                   pending_tasks += 1
                   scope_completion[scope_name]['pending'] += 1

       # Calculate scope percentages
       scope_percentages = {
           scope: round((stats['completed'] / stats['total'] * 100), 2) 
           if stats['total'] > 0 else 0
           for scope, stats in scope_completion.items()
       }

       # Get recent activities
       recent_activities = []
       for execution in execution_data:
           for task in execution.get('tasks', []):
               if 'updated_at' in task:
                   recent_activities.append({
                       'scope_area': execution['scope_area'],
                       'task_name': task['task_name'],
                       'status': task['status'],
                       'team_member': task.get('allocated_team_member', 'Unassigned'),
                       'updated_at': task['updated_at']
                   })

       recent_activities.sort(key=lambda x: x['updated_at'], reverse=True)
       recent_activities = recent_activities[:10]

       # Render PDF template
       env = Environment(loader=FileSystemLoader('templates'))
       template = env.get_template('dashboard_pdf.html')

       html_content = template.render(
           client=client,
           start_date=start_date,
           end_date=end_date,
           total_tasks=total_tasks,
           completed_tasks=completed_tasks,
           in_progress_tasks=in_progress_tasks,
           pending_tasks=pending_tasks,
           scope_percentages=scope_percentages,
           recent_activities=recent_activities
       )

       # Generate PDF using WeasyPrint
       pdf = HTML(string=html_content).write_pdf()
       
       # Send the file
       return send_file(
           io.BytesIO(pdf),
           mimetype='application/pdf',
           as_attachment=True,
           download_name=f'audit_dashboard_{client["company_name"]}_{start_date}_{end_date}.pdf'
       )

   except Exception as e:
       logging.error(f"Error exporting PDF: {str(e)}")
       return jsonify({"error": str(e)}), 500
   
from urllib.parse import unquote
@app.route('/client/<client_id>/annexures', methods=['GET', 'POST'])
def annexures(client_id):
    try:
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')

        if not start_date or not end_date:
            flash("Please select an audit period first", "warning")
            return redirect(url_for('select_period', client_id=client_id))

        # Fetch client data
        client = clients_collection.find_one({"_id": ObjectId(client_id)})
        if not client:
            flash("Client not found", "error")
            return redirect(url_for("dashboard"))

        # Get audit plan
        audit_plan = client_audit_plan_collection.find_one({
            "client_id": ObjectId(client_id),
            "start_date": start_date,
            "end_date": end_date
        })

        if request.method == 'POST':
            try:
                # Get form data
                scope_area = request.form.get('scope_area')
                annexure_name = request.form.get('annexure_name')
                description = request.form.get('description')
                audit_procedure_id = request.form.get('audit_procedure_id')
                table_data = request.form.getlist('table_data[]')

                # Validate required fields
                if not all([scope_area, annexure_name, description]):
                    return jsonify({
                        "success": False,
                        "error": "Please fill in all required fields"
                    }), 400

                # Create annexure document
                annexure_data = {
                    "client_id": ObjectId(client_id),
                    "scope_area": scope_area,
                    "annexure_name": annexure_name,
                    "description": description,
                    "period_start": start_date,
                    "period_end": end_date,
                    "table_data": table_data,
                    "created_at": datetime.now(),
                    "updated_at": datetime.now()
                }

                # Add audit procedure if provided
                if audit_procedure_id:
                    try:
                        annexure_data["audit_procedure_id"] = ObjectId(audit_procedure_id)
                    except:
                        logging.warning(f"Invalid audit procedure ID: {audit_procedure_id}")

                # Insert into database
                result = db.annexures.insert_one(annexure_data)

                if result.inserted_id:
                    return jsonify({
                        "success": True,
                        "message": "Annexure saved successfully",
                        "annexure_id": str(result.inserted_id)
                    })
                else:
                    return jsonify({
                        "success": False,
                        "error": "Failed to save annexure"
                    }), 500

            except Exception as e:
                logging.error(f"Error saving annexure: {str(e)}")
                return jsonify({
                    "success": False,
                    "error": str(e)
                }), 500

        # GET request - fetch existing annexures
        annexures = list(db.annexures.find({
            "client_id": ObjectId(client_id),
            "period_start": start_date,
            "period_end": end_date
        }).sort("created_at", -1))

        # Convert ObjectIds to strings
        client['_id'] = str(client['_id'])
        for annexure in annexures:
            annexure['_id'] = str(annexure['_id'])
            if 'audit_procedure_id' in annexure:
                annexure['audit_procedure_id'] = str(annexure['audit_procedure_id'])

        return render_template(
            'annexures.html',
            client=client,
            audit_plan=audit_plan,
            annexures=annexures,
            start_date=start_date,
            end_date=end_date
        )

    except Exception as e:
        logging.error(f"Error in annexures route: {str(e)}")
        if request.is_xhr or request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return jsonify({
                "success": False,
                "error": str(e)
            }), 500
        flash(f"An error occurred: {str(e)}", "error")
        return redirect(url_for('dashboard'))
# Add these new routes in app.py
from urllib.parse import unquote
@app.route('/client/<client_id>/annexures/view/<annexure_id>', methods=['GET'])
def view_annexure(client_id, annexure_id):
    try:
        # Find the annexure
        annexure = db.annexures.find_one({"_id": ObjectId(annexure_id)})
        if not annexure:
            return jsonify({"success": False, "error": "Annexure not found"}), 404

        # Convert ObjectId to string for JSON serialization
        annexure['_id'] = str(annexure['_id'])
        annexure['client_id'] = str(annexure['client_id'])
        
        # Get audit procedure details if exists
        if 'audit_procedure_id' in annexure:
            try:
                annexure['audit_procedure_id'] = str(annexure['audit_procedure_id'])
                procedure = get_procedure_details(annexure['audit_procedure_id'])
                if procedure:
                    procedure['_id'] = str(procedure['_id'])
                    annexure['procedure_details'] = procedure
            except:
                annexure['procedure_details'] = None

        # Helper function to convert dates to string format
        def format_datetime(obj):
            if isinstance(obj, datetime):
                return obj.strftime('%Y-%m-%d %H:%M:%S')
            return obj

        # Convert datetime objects to strings
        annexure['created_at'] = format_datetime(annexure.get('created_at'))
        annexure['updated_at'] = format_datetime(annexure.get('updated_at'))

        return jsonify({
            "success": True,
            "annexure": annexure
        })

    except Exception as e:
        logging.error(f"Error viewing annexure: {str(e)}")
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500
from urllib.parse import unquote
@app.route('/client/<client_id>/annexures/delete/<annexure_id>', methods=['DELETE'])
def delete_annexure(client_id, annexure_id):
    try:
        # Check if the annexure exists and belongs to the client
        annexure = db.annexures.find_one({
            "_id": ObjectId(annexure_id),
            "client_id": ObjectId(client_id)
        })

        if not annexure:
            return jsonify({
                "success": False,
                "error": "Annexure not found"
            }), 404

        # Delete the annexure
        result = db.annexures.delete_one({
            "_id": ObjectId(annexure_id),
            "client_id": ObjectId(client_id)
        })

        if result.deleted_count > 0:
            return jsonify({
                "success": True,
                "message": "Annexure deleted successfully"
            })
        else:
            return jsonify({
                "success": False,
                "error": "Failed to delete annexure"
            }), 500

    except Exception as e:
        logging.error(f"Error deleting annexure: {str(e)}")
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500
@app.route('/client/<client_id>/annexures/export/<annexure_id>/excel')
def export_annexure_excel(client_id, annexure_id):
    try:
        # Fetch annexure data
        annexure = db.annexures.find_one({"_id": ObjectId(annexure_id)})
        if not annexure:
            flash("Annexure not found", "error")
            return redirect(url_for('annexures', client_id=client_id))

        # Create Excel workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Annexure Details"

        # Add metadata
        ws['A1'] = f"Annexure: {annexure.get('annexure_name', '')}"
        ws['A2'] = f"Scope Area: {annexure.get('scope_area', '')}"
        ws['A3'] = f"Description: {annexure.get('description', '')}"
        ws['A4'] = f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}"

        # Style metadata
        for i in range(1, 5):
            ws[f'A{i}'].font = Font(bold=True)

        # Add empty row for spacing
        current_row = 6

        # Add headers
        headers = ['Sr. No.', 'Document Reference', 'Description', 'Amount', 'Remarks']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=current_row, column=col)
            cell.value = header
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
            cell.alignment = Alignment(horizontal="center", vertical="center")

        # Process and add data rows
        table_data = annexure.get('table_data', [])
        row_data = []
        
        # Group data into rows (4 columns per row)
        for i in range(0, len(table_data), 4):
            row = table_data[i:i+4]
            if len(row) == 4:  # Only add complete rows
                row_data.append(row)

        # Add data rows
        for row_idx, row in enumerate(row_data, 1):
            current_row += 1
            # Add Sr. No.
            ws.cell(row=current_row, column=1, value=row_idx)
            
            # Add rest of the data
            for col_idx, value in enumerate(row, 2):
                cell = ws.cell(row=current_row, column=col_idx)
                try:
                    # Try to convert to number for Amount column
                    if col_idx == 4:  # Amount column
                        cell.value = float(value)
                        cell.number_format = '#,##0.00'
                    else:
                        cell.value = value
                except:
                    cell.value = value
                cell.alignment = Alignment(horizontal="left", vertical="center")

        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = get_column_letter(column[0].column)
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = (max_length + 2) * 1.2
            ws.column_dimensions[column_letter].width = min(adjusted_width, 50)

        # Add borders
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )

        for row in ws.iter_rows(min_row=6, max_row=ws.max_row, min_col=1, max_col=len(headers)):
            for cell in row:
                cell.border = thin_border

        # Save to BytesIO
        excel_file = io.BytesIO()
        wb.save(excel_file)
        excel_file.seek(0)

        return send_file(
            excel_file,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'Annexure_{annexure["annexure_name"]}_{datetime.now().strftime("%Y%m%d")}.xlsx'
        )

    except Exception as e:
        logging.error(f"Error exporting to Excel: {str(e)}")
        flash("Error exporting to Excel", "error")
        return redirect(url_for('annexures', client_id=client_id))

@app.route('/client/<client_id>/annexures/export/<annexure_id>/table')
def export_annexure_table(client_id, annexure_id):
    try:
        # Fetch annexure data
        annexure = db.annexures.find_one({"_id": ObjectId(annexure_id)})
        if not annexure:
            return jsonify({"error": "Annexure not found"}), 404

        # Create Word document
        doc = Document()
        
        # Add title with styling
        title = doc.add_heading(level=1)
        title.add_run('Annexure Report').bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add metadata with styling
        doc.add_paragraph(f'Annexure Name: {annexure.get("annexure_name", "")}').bold = True
        doc.add_paragraph(f'Scope Area: {annexure.get("scope_area", "")}')
        doc.add_paragraph(f'Description: {annexure.get("description", "")}')
        doc.add_paragraph(f'Generated Date: {datetime.now().strftime("%Y-%m-%d %H:%M")}')

        # Add spacing
        doc.add_paragraph()

        # Create table
        table_data = annexure.get('table_data', [])
        rows = []
        
        # Group data into rows (4 columns per row)
        for i in range(0, len(table_data), 4):
            row = table_data[i:i+4]
            if len(row) == 4:  # Only add complete rows
                rows.append(row)

        # Create table
        if rows:
            table = doc.add_table(rows=1, cols=5)  # 5 columns including Sr. No.
            table.style = 'Table Grid'

            # Add headers
            header_cells = table.rows[0].cells
            headers = ['Sr. No.', 'Document Reference', 'Description', 'Amount', 'Remarks']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].font.bold = True
                header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add data rows
            for idx, row_data in enumerate(rows, 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(idx)  # Sr. No.
                
                for col, value in enumerate(row_data, 1):
                    row_cells[col].text = str(value)
                    row_cells[col].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Set column widths
            widths = [Inches(0.8), Inches(2), Inches(2), Inches(1.2), Inches(2)]
            for row in table.rows:
                for idx, cell in enumerate(row.cells):
                    cell._tc.tcPr.tcW.type = 'dxa'
                    cell._tc.tcPr.tcW.w = int(widths[idx].inches * 1440)

        # Save to BytesIO
        doc_file = io.BytesIO()
        doc.save(doc_file)
        doc_file.seek(0)

        return send_file(
            doc_file,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'Annexure_{annexure["annexure_name"]}_{datetime.now().strftime("%Y%m%d")}.docx'
        )

    except Exception as e:
        logging.error(f"Error exporting to Word: {str(e)}")
        flash("Error exporting to Word", "error")
        return redirect(url_for('annexures', client_id=client_id))
    
@app.route('/client/<client_id>/annexures/download-template')
def download_annexure_template(client_id):
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "Annexure Data"

        # Define headers
        headers = ['Sr. No.', 'Document Reference', 'Description', 'Amount', 'Remarks']

        # Style for headers
        header_fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
        header_font = Font(bold=True)
        header_alignment = Alignment(horizontal='center', vertical='center')

        # Add headers
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col)
            cell.value = header
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = header_alignment

        # Add sample data
        sample_data = [
            [1, "INV-001", "Sample Invoice", 1000.00, "Sample Remarks"],
            [2, "DOC-002", "Sample Document", 2000.00, "Additional Notes"]
        ]

        for row_idx, data in enumerate(sample_data, 2):
            for col_idx, value in enumerate(data, 1):
                cell = ws.cell(row=row_idx, column=col_idx)
                cell.value = value
                cell.font = Font(color="808080")

        # Set column widths
        column_widths = [10, 20, 40, 15, 30]
        for i, width in enumerate(column_widths):
            ws.column_dimensions[get_column_letter(i + 1)].width = width

        # Add instructions sheet
        ws_instructions = wb.create_sheet("Instructions")
        instructions = [
            ["Template Instructions:"],
            ["1. Do not modify the column headers"],
            ["2. Document Reference should be unique"],
            ["3. Amount should be numeric without currency symbols"],
            ["4. Maximum 1000 rows allowed"],
            ["5. Remove the sample rows before adding actual data"],
            ["6. Save as .xlsx format before uploading"],
            ["7. All columns are required"]
        ]

        for row_idx, instruction in enumerate(instructions, 1):
            cell = ws_instructions.cell(row=row_idx, column=1, value=instruction[0])
            if row_idx == 1:
                cell.font = Font(bold=True)

        # Save to BytesIO
        excel_file = io.BytesIO()
        wb.save(excel_file)
        excel_file.seek(0)

        return send_file(
            excel_file,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name='annexure_template.xlsx'
        )

    except Exception as e:
        logging.error(f"Error creating template: {str(e)}")
        return jsonify({"error": "Failed to generate template"}), 500

@app.route('/client/<client_id>/annexures/import', methods=['POST'])
def import_annexure(client_id):
    try:
        if 'file' not in request.files:
            return jsonify({"success": False, "error": "No file uploaded"}), 400

        file = request.files['file']
        if not file or file.filename == '':
            return jsonify({"success": False, "error": "No file selected"}), 400

        if not file.filename.endswith('.xlsx'):
            return jsonify({"success": False, "error": "Only .xlsx files are allowed"}), 400

        # Get form data
        scope_area = request.form.get('scope_area')
        audit_procedure = request.form.get('audit_procedure')
        annexure_name = request.form.get('annexure_name')
        description = request.form.get('description')
        start_date = request.form.get('start_date')
        end_date = request.form.get('end_date')

        # Validate required fields
        if not all([scope_area, annexure_name, description, start_date, end_date]):
            return jsonify({
                "success": False,
                "error": "Missing required fields"
            }), 400

        try:
            # Read Excel file
            df = pd.read_excel(file)
            if df.empty:
                return jsonify({
                    "success": False,
                    "error": "Excel file is empty"
                }), 400

            # Validate Excel structure
            required_columns = ['Document Reference', 'Description', 'Amount', 'Remarks']
            missing_columns = [col for col in required_columns if col not in df.columns]
            if missing_columns:
                return jsonify({
                    "success": False,
                    "error": f"Missing columns in Excel: {', '.join(missing_columns)}"
                }), 400

            # Process Excel data
            table_data = []
            for _, row in df.iterrows():
                # Convert each cell to string and handle NaN/None values
                doc_ref = str(row['Document Reference']) if pd.notna(row['Document Reference']) else ''
                desc = str(row['Description']) if pd.notna(row['Description']) else ''
                amount = str(row['Amount']) if pd.notna(row['Amount']) else '0'
                remarks = str(row['Remarks']) if pd.notna(row['Remarks']) else ''
                
                table_data.extend([doc_ref, desc, amount, remarks])

            # Create annexure document
            annexure_data = {
                "client_id": ObjectId(client_id),
                "scope_area": scope_area,
                "annexure_name": annexure_name,
                "description": description,
                "period_start": start_date,
                "period_end": end_date,
                "table_data": table_data,
                "created_at": datetime.now(),
                "updated_at": datetime.now()
            }

            # Add audit procedure if provided
            if audit_procedure:
                try:
                    annexure_data["audit_procedure_id"] = ObjectId(audit_procedure)
                except:
                    logging.warning(f"Invalid audit procedure ID: {audit_procedure}")

            # Insert into database
            result = db.annexures.insert_one(annexure_data)

            if result.inserted_id:
                return jsonify({
                    "success": True,
                    "message": "Annexure imported successfully",
                    "annexure_id": str(result.inserted_id)
                })
            else:
                return jsonify({
                    "success": False,
                    "error": "Failed to save annexure"
                }), 500

        except Exception as e:
            logging.error(f"Excel processing error: {str(e)}")
            return jsonify({
                "success": False,
                "error": f"Error processing Excel file: {str(e)}"
            }), 500

    except Exception as e:
        logging.error(f"Import error: {str(e)}")
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500
def validate_excel_file(file_stream):
    """Validate Excel file structure and content"""
    try:
        df = pd.read_excel(file_stream)
        required_columns = ['Document Reference', 'Description', 'Amount', 'Remarks']
        
        # Check if all required columns exist (ignoring case)
        df_columns = [col.strip().lower() for col in df.columns]
        missing_columns = [col for col in required_columns 
                         if col.lower() not in df_columns]
        
        if missing_columns:
            return False, f"Missing required columns: {', '.join(missing_columns)}"
            
        return True, df
    except Exception as e:
        return False, f"Invalid Excel file: {str(e)}"


@app.route('/client/<client_id>/annexures/edit/<annexure_id>', methods=['GET', 'POST'])
def edit_annexure(client_id, annexure_id):
    try:
        # Get period dates and annexure data
        annexure = db.annexures.find_one({"_id": ObjectId(annexure_id)})
        if not annexure:
            flash("Annexure not found", "error")
            return redirect(url_for('annexures', client_id=client_id))

        start_date = request.args.get('start_date') or annexure.get('period_start')
        end_date = request.args.get('end_date') or annexure.get('period_end')

        # Fetch client
        client = clients_collection.find_one({"_id": ObjectId(client_id)})
        if not client:
            flash("Client not found", "error")
            return redirect(url_for('dashboard'))

        # Fetch audit plan
        audit_plan = client_audit_plan_collection.find_one({
            "client_id": ObjectId(client_id),
            "start_date": start_date,
            "end_date": end_date
        })

        if request.method == 'POST':
            try:
                # Get form data
                scope_area = request.form.get('scope_area')
                audit_procedure_id = request.form.get('audit_procedure_id')
                annexure_name = request.form.get('annexure_name')
                description = request.form.get('description')
                raw_table_data = request.form.getlist('table_data[]')

                # Process table data into rows
                processed_table_data = []
                for i in range(0, len(raw_table_data), 4):
                    if i + 4 <= len(raw_table_data):
                        row_data = raw_table_data[i:i+4]
                        processed_table_data.extend(row_data)

                # Build base update document
                update_doc = {
                    "$set": {
                        "scope_area": scope_area,
                        "annexure_name": annexure_name,
                        "description": description,
                        "table_data": processed_table_data,
                        "updated_at": datetime.now(),
                        "period_start": start_date,
                        "period_end": end_date
                    }
                }

                # Handle audit procedure ID separately
                if audit_procedure_id:
                    update_doc["$set"]["audit_procedure_id"] = ObjectId(audit_procedure_id)
                elif 'audit_procedure_id' in annexure:
                    update_doc["$unset"] = {"audit_procedure_id": ""}

                # Update document in MongoDB
                update_result = db.annexures.update_one(
                    {"_id": ObjectId(annexure_id)},
                    update_doc
                )

                if update_result.modified_count > 0:
                    flash("Annexure updated successfully!", "success")
                else:
                    flash("No changes were made", "info")

                return redirect(url_for('annexures', 
                                      client_id=client_id,
                                      start_date=start_date,
                                      end_date=end_date))

            except Exception as e:
                logging.error(f"Error updating annexure: {str(e)}")
                flash(f"Error updating annexure: {str(e)}", "error")

        # Convert ObjectIds to strings for template
        client['_id'] = str(client['_id'])
        annexure['_id'] = str(annexure['_id'])
        
        # Handle audit_procedure_id for template
        if 'audit_procedure_id' in annexure:
            annexure['audit_procedure_id'] = str(annexure['audit_procedure_id'])
            
            # Get procedure details if available
            procedure = get_procedure_details(annexure['audit_procedure_id'])
            if procedure:
                annexure['procedure_details'] = procedure

        # Get all scope areas from audit plan
        scope_areas = []
        if audit_plan and 'modified_audit_scope' in audit_plan:
            scope_areas = [{'name': scope['name']} for scope in audit_plan['modified_audit_scope']]

        # Ensure current scope area is included
        current_scope = {'name': annexure['scope_area']}
        if current_scope not in scope_areas:
            scope_areas.append(current_scope)

        # Update audit plan for template
        if not audit_plan:
            audit_plan = {}
        audit_plan['modified_audit_scope'] = scope_areas

        return render_template(
            'edit_annexure.html',
            client=client,
            annexure=annexure,
            audit_plan=audit_plan,
            start_date=start_date,
            end_date=end_date
        )

    except Exception as e:
        logging.error(f"Error in edit_annexure: {str(e)}")
        flash(f"Error processing the request: {str(e)}", "error")
        return redirect(url_for('annexures', 
                              client_id=client_id,
                              start_date=start_date,
                              end_date=end_date))

# Update the annexure creation route
@app.route('/client/<client_id>/annexures', methods=['POST'])
def create_annexure(client_id):
    try:
        # Log incoming form data
        logging.info("Received form data:")
        for key, value in request.form.items():
            logging.info(f"{key}: {value}")

        # Get audit procedure ID
        audit_procedure_id = request.form.get('audit_procedure_id')
        logging.info(f"Received audit_procedure_id: {audit_procedure_id}")

        # Create annexure data
        annexure_data = {
            "client_id": ObjectId(client_id),
            "period_start": request.form.get('start_date'),
            "period_end": request.form.get('end_date'),
            "scope_area": request.form.get('scope_area'),
            "annexure_name": request.form.get('annexure_name'),
            "description": request.form.get('description'),
            "table_data": request.form.getlist('table_data[]'),
            "created_at": datetime.now(),
            "updated_at": datetime.now()
        }

        # Add audit procedure ID if present
        if audit_procedure_id and audit_procedure_id.strip():
            try:
                annexure_data["audit_procedure_id"] = ObjectId(audit_procedure_id)
                logging.info(f"Added audit_procedure_id to document: {audit_procedure_id}")
            except Exception as e:
                logging.error(f"Error converting audit_procedure_id: {str(e)}")

        # Log final data before insertion
        logging.info(f"Final annexure data: {annexure_data}")

        # Insert into database
        result = db.annexures.insert_one(annexure_data)
        
        if result.inserted_id:
            # Verify the saved document
            saved_doc = db.annexures.find_one({"_id": result.inserted_id})
            logging.info(f"Successfully saved document: {saved_doc}")
            flash("Annexure saved successfully!", "success")
            return jsonify({"success": True, "message": "Annexure saved successfully"})
        else:
            return jsonify({"success": False, "error": "Failed to save annexure"}), 500

    except Exception as e:
        logging.error(f"Error creating annexure: {str(e)}")
        return jsonify({"success": False, "error": str(e)}), 500
    
@app.route('/api/verify-procedure/<procedure_id>')
def verify_procedure(procedure_id):
    try:
        procedure = audit_execution_dynamic.find_one(
            {"tasks._id": ObjectId(procedure_id)},
            {"tasks.$": 1}
        )
        return jsonify({
            "exists": bool(procedure),
            "details": procedure['tasks'][0] if procedure else None
        })
    except Exception as e:
        logging.error(f"Error verifying procedure: {str(e)}")
        return jsonify({"error": str(e)}), 500
    
@app.route('/client/<client_id>/get-audit-procedures/<scope_area>')
def get_audit_procedures(client_id, scope_area):
    try:
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')

        if not all([start_date, end_date]):
            return jsonify({
                "success": False,
                "error": "Missing date parameters"
            }), 400

        # Fetch procedures from audit_execution_dynamic for the specific period
        execution_data = list(audit_execution_dynamic.find({
            "client_id": ObjectId(client_id),
            "scope_area": scope_area,
            "period_start": start_date,
            "period_end": end_date,
            "tasks": {"$exists": True, "$not": {"$size": 0}}
        }))

        procedures = []
        if execution_data:
            for execution in execution_data:
                for task in execution.get('tasks', []):
                    procedures.append({
                        "id": str(task.get('_id')),
                        "name": task.get('task_name'),
                        "procedure": task.get('procedure')
                    })

        return jsonify({
            "success": True,
            "procedures": procedures
        })

    except Exception as e:
        logging.error(f"Error fetching audit procedures: {str(e)}")
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500
# Add helper function to get procedure details
def get_procedure_details(procedure_id):
    try:
        # First try to find in audit_execution_dynamic
        task = audit_execution_dynamic.find_one(
            {"tasks._id": ObjectId(procedure_id)},
            {"tasks.$": 1}
        )
        
        if task and task.get('tasks'):
            return task['tasks'][0]
            
        return None
    except Exception as e:
        logging.error(f"Error getting procedure details: {str(e)}")
        return None


def validate_excel_data(df):
    """Validate Excel file data structure and content"""
    try:
        required_columns = ['Document Reference', 'Description', 'Amount']
        
        # Check required columns
        if not all(col in df.columns for col in required_columns):
            return False, "Missing required columns"
            
        # Check data types
        if not pd.to_numeric(df['Amount'], errors='coerce').notnull().all():
            return False, "Amount column contains non-numeric values"
            
        # Check empty values
        if df[required_columns].isnull().any().any():
            return False, "Required columns contain empty values"
            
        return True, df
    except Exception as e:
        return False, f"Invalid data format: {str(e)}"


# Register the helper function with Jinja2
app.jinja_env.globals.update(get_procedure_details=get_procedure_details)
@app.route('/client/<client_id>/annexures/view/<annexure_id>', methods=['GET'])
def view_annexure_details(client_id, annexure_id):
    try:
        annexure = db.annexures.find_one({"_id": ObjectId(annexure_id)})
        if not annexure:
            return jsonify({"success": False, "error": "Annexure not found"}), 404

        # If there's a linked procedure, get its details
        if annexure.get('audit_procedure_id'):
            procedure = get_procedure_details(annexure['audit_procedure_id'])
            annexure['procedure_name'] = procedure.get('task_name') if procedure else None

        # Convert ObjectId to string for JSON serialization
        annexure['_id'] = str(annexure['_id'])
        if 'client_id' in annexure:
            annexure['client_id'] = str(annexure['client_id'])

        return jsonify({
            "success": True,
            "annexure": annexure
        })

    except Exception as e:
        logging.error(f"Error viewing annexure: {str(e)}")
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500

@app.route('/client/<client_id>/annexures', methods=['POST'])
def save_annexure(client_id):
    try:
        data = request.form
        
        # Validate required fields
        required_fields = ['scope_area', 'annexure_name', 'description']
        for field in required_fields:
            if not data.get(field):
                return jsonify({
                    "success": False,
                    "error": f"Missing required field: {field}"
                }), 400

        annexure_data = {
            "client_id": ObjectId(client_id),
            "scope_area": data.get('scope_area'),
            "audit_procedure_id": ObjectId(data.get('audit_procedure')) if data.get('audit_procedure') else None,
            "annexure_name": data.get('annexure_name'),
            "description": data.get('description'),
            "table_data": request.form.getlist('table_data[]'),
            "created_at": datetime.now(),
            "updated_at": datetime.now()
        }

        result = db.annexures.insert_one(annexure_data)

        if result.inserted_id:
            return jsonify({
                "success": True,
                "message": "Annexure saved successfully"
            })
        else:
            return jsonify({
                "success": False,
                "error": "Failed to save annexure"
            }), 500

    except Exception as e:
        logging.error(f"Error saving annexure: {str(e)}")
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/client/<client_id>/get_available_annexures', methods=['GET'])
def get_available_annexures(client_id):
    try:
        task_id = request.args.get('task_id')
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')
        
        # Get all annexures for the client and period
        annexures = list(db.annexures.find({
            "client_id": ObjectId(client_id),
            "period_start": start_date,
            "period_end": end_date
        }))
        
        # Get current task details for checking existing links
        task = audit_execution_dynamic.find_one(
            {"tasks._id": ObjectId(task_id)},
            {"tasks.$": 1}
        )
        
        # Get IDs of already linked annexures
        linked_annexure_ids = []
        if task and 'tasks' in task and len(task['tasks']) > 0:
            linked_annexure_ids = [str(annexure['id']) for annexure in task['tasks'][0].get('linked_annexures', [])]
        
        # Format response data
        formatted_annexures = []
        for annexure in annexures:
            formatted_annexures.append({
                "id": str(annexure['_id']),
                "name": annexure.get('annexure_name', ''),
                "description": annexure.get('description', ''),
                "created_at": annexure['created_at'].strftime("%Y-%m-%d %H:%M"),
                "is_linked": str(annexure['_id']) in linked_annexure_ids
            })
        
        return jsonify({
            "success": True,
            "annexures": formatted_annexures
        })
        
    except Exception as e:
        logging.error(f"Error getting available annexures: {str(e)}")
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/client/<client_id>/link_annexures/<task_id>', methods=['POST'])
def link_annexures(client_id, task_id):
    try:
        data = request.get_json()
        annexure_ids = data.get('annexure_ids', [])
        
        # Convert string IDs to ObjectId
        annexure_ids = [ObjectId(id) for id in annexure_ids]
        
        # Get annexure details
        annexures = list(db.annexures.find({
            "_id": {"$in": annexure_ids}
        }, {
            "_id": 1,
            "annexure_name": 1
        }))
        
        # Format linked annexures data
        linked_annexures = [
            {
                "id": str(annexure['_id']),
                "name": annexure['annexure_name']
            }
            for annexure in annexures
        ]
        
        # Update task with linked annexures
        result = audit_execution_dynamic.update_one(
            {"tasks._id": ObjectId(task_id)},
            {
                "$set": {
                    "tasks.$.linked_annexures": linked_annexures,
                    "tasks.$.updated_at": datetime.now()
                }
            }
        )
        
        return jsonify({
            "success": True,
            "message": "Annexures linked successfully",
            "linked_annexures": linked_annexures
        })
        
    except Exception as e:
        logging.error(f"Error linking annexures: {str(e)}")
        return jsonify({"success": False, "error": str(e)}), 500

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

@app.route('/client/<client_id>/export-company-info')
def export_company_info(client_id):
    try:
        # Fetch client data
        client = clients_collection.find_one({"_id": ObjectId(client_id)})
        if not client:
            flash('Client not found', 'error')
            return redirect(url_for('dashboard'))

        # Create document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Company Information', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add company name as subtitle
        company_name = doc.add_heading(client['company_name'], level=1)
        company_name.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Function to add section
        def add_section(title, data_dict, keys):
            doc.add_heading(title, level=2)
            table = doc.add_table(rows=0, cols=2)
            table.style = 'Table Grid'
            for key in keys:
                row = table.add_row()
                row.cells[0].text = key.replace('_', ' ').title() + ':'
                row.cells[1].text = str(data_dict.get(key, 'N/A'))
            doc.add_paragraph()  # Add spacing after table

        # Basic Information
        basic_info_keys = [
            'business_registration_number', 'gstin', 'pan_tan', 
            'registered_address', 'operating_addresses', 'key_management_personnel',
            'ownership_structure', 'date_of_incorporation', 'authorized_signatories',
            'bank_account_details'
        ]
        add_section('Basic Information', client, basic_info_keys)

        # Industry and Business Model
        industry_keys = [
            'industry_type', 'other_industry_type', 'nature_of_products',
            'key_raw_materials', 'key_markets', 'major_competitors',
            'manufacturing_processes', 'key_technology', 'capacity_utilization',
            'regulatory_requirements'
        ]
        add_section('Industry and Business Model', client, industry_keys)

        # Products and Services
        products_keys = [
            'product_categories', 'top_selling_products', 'production_capacity',
            'lead_time', 'product_differentiation', 'product_lifecycle',
            'pricing_model', 'after_sales_services'
        ]
        add_section('Products and Services', client, products_keys)

        # Ownership and Governance
        governance_keys = [
            'ownership_structure_details', 'board_of_directors',
            'corporate_governance', 'legal_advisors_auditors', 'key_policies'
        ]
        add_section('Ownership and Governance', client, governance_keys)

        # Financial Overview
        financial_keys = [
            'revenue_streams', 'cost_structure', 'profit_margins',
            'capital_structure', 'financial_ratios', 'tax_compliance'
        ]
        add_section('Financial Overview', client, financial_keys)

        # HR and Payroll
        hr_keys = [
            'employee_count', 'employee_turnover_rate', 'compensation_benefits',
            'training_programs', 'performance_management', 'workforce_diversity',
            'labor_union'
        ]
        add_section('Human Resources and Payroll', client, hr_keys)

        # IT Systems
        it_keys = [
            'erp_software', 'data_security_measures', 'backup_disaster_recovery',
            'system_integrations', 'user_access_control'
        ]
        add_section('IT Systems and Infrastructure', client, it_keys)

        # Risk Management
        risk_keys = [
            'key_risks', 'regulatory_compliance', 'insurance_policies',
            'internal_controls'
        ]
        add_section('Risk Management and Compliance', client, risk_keys)

        # Save to memory buffer
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        return send_file(
            file_stream,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'Company_Information_{client["company_name"]}.docx'
        )

    except Exception as e:
        flash(f'Error generating document: {str(e)}', 'error')
        return redirect(url_for('company_information', client_id=client_id))

@app.route('/get_dynamic_audit_scope', methods=['GET'])
def get_dynamic_audit_scope():
    try:
        # Fetch all scope items from dynamic_audit_scope collection
        scope_items = list(db.dynamic_audit_scope.find())
        
        # Convert ObjectIds to strings for JSON serialization
        for item in scope_items:
            item['_id'] = str(item['_id'])
        
        return jsonify({
            "success": True,
            "scope_items": scope_items
        })
    except Exception as e:
        logging.error(f"Error fetching dynamic audit scope: {str(e)}")
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500
    
# Change this route in app.py
@app.route('/client/<client_id>/overview')
def client_overview(client_id):
    try:
        # Get client information
        client = clients_collection.find_one({"_id": ObjectId(client_id)})
        if not client:
            flash("Client not found", "error")
            return redirect(url_for("dashboard"))
            
        # Get audit statistics
        audit_stats = get_audit_statistics(client_id)
        
        # Get recent audit executions
        audit_executions = list(audit_execution_dynamic.find({
            "client_id": ObjectId(client_id)
        }).sort("updated_at", -1).limit(5))  # Get 5 most recent executions
        
        return render_template(
            'overview.html',
            client=client,
            audit_stats=audit_stats,
            audit_executions=audit_executions
        )
        
    except Exception as e:
        logging.error(f"Error loading overview page: {str(e)}")
        flash("An error occurred while loading the overview page", "error")
        return redirect(url_for("dashboard"))

def get_audit_statistics(client_id):
    """Get audit progress statistics for a client"""
    try:
        stats = {
            "total_tasks": 0,
            "completed_tasks": 0,
            "in_progress": 0,
            "pending": 0
        }
        
        executions = audit_execution_dynamic.find({
            "client_id": ObjectId(client_id)
        })
        
        for execution in executions:
            for task in execution.get('tasks', []):
                stats['total_tasks'] += 1
                status = task.get('status', '').lower()
                if status == 'completed':
                    stats['completed_tasks'] += 1
                elif status == 'in progress':
                    stats['in_progress'] += 1
                else:
                    stats['pending'] += 1
                    
        return stats
        
    except Exception as e:
        logging.error(f"Error getting audit statistics: {str(e)}")
        return None


@app.route('/client/<client_id>/export_overview')
def export_overview(client_id):
    try:
        # Get client information
        client = clients_collection.find_one({"_id": ObjectId(client_id)})
        if not client:
            flash("Client not found", "error")
            return redirect(url_for("dashboard"))

        # Get audit statistics
        audit_stats = get_audit_statistics(client_id)
        
        # Get recent audit executions
        audit_executions = list(audit_execution_dynamic.find({
            "client_id": ObjectId(client_id)
        }).sort("updated_at", -1).limit(5))

        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'Company Overview - {client["company_name"]}', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
        
        # Add Quick Stats
        doc.add_heading('Quick Statistics', level=2)
        stats_table = doc.add_table(rows=1, cols=2)
        stats_table.style = 'Table Grid'
        stats_cells = stats_table.rows[0].cells
        stats_cells[0].text = 'Metric'
        stats_cells[1].text = 'Value'
        
        stats_data = [
            ('Total Revenue', client.get('revenue_streams', 'N/A')),
            ('Employee Count', client.get('employee_count', 'N/A')),
            ('Completed Tasks', str(audit_stats.get('completed_tasks', 0))),
            ('Pending Tasks', str(audit_stats.get('pending', 0)))
        ]
        
        for metric, value in stats_data:
            row = stats_table.add_row()
            row.cells[0].text = metric
            row.cells[1].text = str(value)

        # Add Company Information
        doc.add_heading('Company Information', level=2)
        company_table = doc.add_table(rows=1, cols=2)
        company_table.style = 'Table Grid'
        header_cells = company_table.rows[0].cells
        header_cells[0].text = 'Field'
        header_cells[1].text = 'Details'
        
        company_data = [
            ('Registration Number', client.get('business_registration_number', 'N/A')),
            ('GSTIN', client.get('gstin', 'N/A')),
            ('PAN/TAN', client.get('pan_tan', 'N/A')),
            ('Incorporation Date', client.get('date_of_incorporation', 'N/A')),
            ('Key Management', client.get('key_management_personnel', 'N/A'))
        ]
        
        for field, value in company_data:
            row = company_table.add_row()
            row.cells[0].text = field
            row.cells[1].text = str(value)

        # Add Recent Audit Items
        doc.add_heading('Recent Audit Items', level=2)
        if audit_executions:
            audit_table = doc.add_table(rows=1, cols=3)
            audit_table.style = 'Table Grid'
            header_cells = audit_table.rows[0].cells
            header_cells[0].text = 'Scope Area'
            header_cells[1].text = 'Task'
            header_cells[2].text = 'Status'
            
            for execution in audit_executions:
                for task in execution.get('tasks', [])[:3]:  # Show only first 3 tasks
                    row = audit_table.add_row()
                    row.cells[0].text = execution.get('scope_area', 'N/A')
                    row.cells[1].text = task.get('task_name', 'N/A')
                    row.cells[2].text = task.get('status', 'N/A')
        else:
            doc.add_paragraph('No recent audit items found')

        # Add Risk Assessment
        doc.add_heading('Risk Assessment', level=2)
        doc.add_paragraph(f'Key Risks: {client.get("key_risks", "No risks identified")}')
        doc.add_paragraph(f'Internal Controls: {client.get("internal_controls", "No internal controls specified")}')
        doc.add_paragraph(f'Compliance Status: {client.get("regulatory_compliance", "No compliance information available")}')

        # Save to BytesIO
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)

        return send_file(
            doc_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'Overview_{client["company_name"]}_{datetime.now().strftime("%Y%m%d")}.docx'
        )

    except Exception as e:
        logging.error(f"Error exporting overview: {str(e)}")
        flash("Error generating overview document", "error")
        return redirect(url_for('client_overview', client_id=client_id))

def get_client_storage_info(client_id):
    try:
        # Find all files for this specific client
        files = db['fs.files'].find({"metadata.client_id": client_id})
        
        # Calculate total storage used by this client
        total_size = sum(file.get('length', 0) for file in files)
        
        # Count total number of files for this client
        file_count = db['fs.files'].count_documents({"metadata.client_id": client_id})
        
        # Convert bytes to readable format
        def convert_size(size_bytes):
            for unit in ['B', 'KB', 'MB', 'GB']:
                if size_bytes < 1024.0:
                    return f"{size_bytes:.2f} {unit}"
                size_bytes /= 1024.0
            return f"{size_bytes:.2f} TB"

        return {
            'total_size': convert_size(total_size),
            'file_count': file_count,
            'raw_size': total_size
        }
    except Exception as e:
        logging.error(f"Error calculating client storage info: {str(e)}")
        return {
            'total_size': '0 MB',
            'file_count': 0,
            'raw_size': 0
        }
def check_storage_limit(client_id, new_file_size):
    """Check if adding new file would exceed storage limit of 50MB for the client"""
    try:
        # Get current storage info for this client
        client_storage = get_client_storage_info(client_id)
        current_size = client_storage['raw_size']
            
        # Convert to MB for readable message
        current_usage_mb = current_size/(1024*1024)
            
        # Check if adding new file would exceed limit
        if (current_size + new_file_size) > STORAGE_LIMIT:
            return False, f"Client storage limit of 50MB exceeded. Current usage: {current_usage_mb:.2f}MB"
        return True, None
        
    except Exception as e:
        logging.error(f"Error checking client storage limit: {str(e)}")
        return False, "Error checking storage capacity"

@app.route('/track_visitor', methods=['POST'])
def track_visitor():
    try:
        visitor_data = request.json
        
        # Parse user agent for better device info
        ua_string = visitor_data.get('user_agent')
        user_agent = parse(ua_string)

        # Enhance visitor data
        enhanced_data = {
            'visitor_id': visitor_data.get('visitor_id'),
            'timestamp': datetime.now(),
            'page_url': visitor_data.get('page_url'),
            'referrer': visitor_data.get('referrer'),
            'screen_resolution': visitor_data.get('screen_resolution'),
            'browser': user_agent.browser.family,
            'browser_version': user_agent.browser.version_string,
            'os': user_agent.os.family,
            'device': user_agent.device.family,
            'is_mobile': user_agent.is_mobile,
            'is_tablet': user_agent.is_tablet,
            'is_pc': user_agent.is_pc,
            'language': visitor_data.get('language'),
            'ip_address': request.remote_addr
        }

        # Store in MongoDB
        mongo.db.visitors.insert_one(enhanced_data)
        
        # Create session if doesn't exist
        if 'visitor_id' not in session:
            session['visitor_id'] = visitor_data.get('visitor_id')

        return jsonify({'status': 'success'}), 200

    except Exception as e:
        logging.error(f"Error tracking visitor: {str(e)}")
        return jsonify({'status': 'error', 'message': str(e)}), 500

# Add route to view visitor analytics (protected)
@app.route('/visitor_analytics')
def visitor_analytics():
    if 'logged_in' not in session:
        return redirect(url_for('login'))
        
    try:
        # Get visitor statistics
        total_visitors = mongo.db.visitors.count_documents({})
        unique_visitors = len(mongo.db.visitors.distinct('visitor_id'))
        
        # Get device breakdown
        device_stats = {
            'mobile': mongo.db.visitors.count_documents({'is_mobile': True}),
            'tablet': mongo.db.visitors.count_documents({'is_tablet': True}),
            'desktop': mongo.db.visitors.count_documents({'is_pc': True})
        }
        
        # Get recent visitors
        recent_visitors = list(mongo.db.visitors.find(
            {},
            {'timestamp': 1, 'browser': 1, 'os': 1, 'device': 1}
        ).sort('timestamp', -1).limit(10))

        return render_template(
            'visitor_analytics.html',
            total_visitors=total_visitors,
            unique_visitors=unique_visitors,
            device_stats=device_stats,
            recent_visitors=recent_visitors
        )

    except Exception as e:
        logging.error(f"Error getting visitor analytics: {str(e)}")
        flash("Error retrieving analytics data", "error")
        return redirect(url_for('dashboard'))
        
@app.route('/client/<client_id>/task/<task_id>/upload-documents', methods=['POST'])
def upload_task_documents(client_id, task_id):
    try:
        if 'documents' not in request.files:
            return jsonify({"success": False, "error": "No files uploaded"}), 400
            
        files = request.files.getlist('documents')
        uploaded_count = 0
        
        for file in files:
            if file.filename == '':
                continue
                
            # Check file size
            file_content = file.read()
            file_size = len(file_content)
            if file_size > 15 * 1024 * 1024:  # 15MB limit
                continue
                
            # Reset file pointer
            file.seek(0)
            
            # Get MIME type
            mime_type = file.content_type or mimetypes.guess_type(file.filename)[0]
            
            # Store file metadata
            doc_metadata = {
                "client_id": client_id,
                "task_id": task_id,
                "title": request.form.get('doc_title', file.filename),
                "description": request.form.get('doc_description', ''),
                "uploaded_at": datetime.now(),
                "uploaded_by": session.get('username'),
                "mime_type": mime_type,
                "size": file_size
            }
            
            # Store in GridFS
            file_id = fs.put(file, 
                           filename=file.filename,
                           metadata=doc_metadata,
                           content_type=mime_type)
            
            # Update task document
            audit_execution_dynamic.update_one(
                {"tasks._id": ObjectId(task_id)},
                {
                    "$push": {
                        "tasks.$.supporting_docs": {
                            "_id": file_id,
                            "filename": file.filename,
                            "title": doc_metadata["title"],
                            "mime_type": mime_type,
                            "size": file_size,
                            "uploaded_at": doc_metadata["uploaded_at"]
                        }
                    }
                }
            )
            
            uploaded_count += 1
        
        return jsonify({
            "success": True,
            "message": f"{uploaded_count} documents uploaded successfully",
            "uploaded_count": uploaded_count
        })
        
    except Exception as e:
        logging.error(f"Error uploading task documents: {str(e)}")
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/client/<client_id>/task/<task_id>/documents')
def get_task_documents(client_id, task_id):
    try:
        task = audit_execution_dynamic.find_one(
            {"tasks._id": ObjectId(task_id)},
            {"tasks.$": 1}
        )
        
        if not task or not task.get('tasks'):
            return jsonify({
                "success": False,
                "error": "Task not found"
            }), 404
            
        supporting_docs = task['tasks'][0].get('supporting_docs', [])
        
        documents = []
        for doc in supporting_docs:
            try:
                file_obj = fs.get(doc['_id'])
                documents.append({
                    "_id": str(doc['_id']),
                    "filename": doc['filename'],
                    "title": doc['title'],
                    "description": file_obj.metadata.get('description', ''),
                    "uploaded_at": doc['uploaded_at'],
                    "mime_type": doc.get('mime_type', 'application/octet-stream'),
                    "size": doc.get('size', 0)
                })
            except Exception as e:
                logging.error(f"Error retrieving document {doc['_id']}: {str(e)}")
                continue
                
        return jsonify({
            "success": True,
            "documents": documents
        })
        
    except Exception as e:
        logging.error(f"Error fetching task documents: {str(e)}")
        return jsonify({"success": False, "error": str(e)}), 500

from werkzeug.utils import secure_filename
from urllib.parse import quote

@app.route('/client/<client_id>/task-document/<doc_id>/download')
def download_task_document(client_id, doc_id):
    try:
        # Convert string ID to ObjectId
        file_id = ObjectId(doc_id)
        
        # Get file from GridFS
        grid_out = fs.get(file_id)
        
        if not grid_out:
            flash('File not found', 'error')
            return redirect(url_for('audit_execution', client_id=client_id))

        # Get the original filename and ensure extension is preserved
        original_filename = grid_out.filename
        # Clean the filename to remove any potentially unsafe characters
        safe_filename = secure_filename(original_filename)

        # Create response with file data
        response = send_file(
            io.BytesIO(grid_out.read()),
            mimetype=grid_out.content_type,
            as_attachment=True,
            download_name=safe_filename  # Use the cleaned original filename
        )
        
        # Set additional headers to force download with correct filename
        response.headers["Content-Disposition"] = f"attachment; filename*=UTF-8''{quote(safe_filename)}"
        response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
        response.headers["Pragma"] = "no-cache"
        response.headers["Expires"] = "0"
        
        return response

    except Exception as e:
        logging.error(f"Error downloading task document: {str(e)}")
        flash('Error downloading file', 'error')
        return redirect(url_for('audit_execution', client_id=client_id))

@app.route('/client/<client_id>/task-document/<doc_id>/delete', methods=['POST'])
def delete_task_document(client_id, doc_id):
    try:
        # Convert string ID to ObjectId
        file_id = ObjectId(doc_id)
        
        # Delete file from GridFS
        fs.delete(file_id)
        
        # Remove document reference from task
        result = audit_execution_dynamic.update_many(
            {"client_id": ObjectId(client_id)},
            {
                "$pull": {
                    "tasks.$[].supporting_docs": {
                        "_id": file_id
                    }
                }
            }
        )
        
        if result.modified_count > 0:
            return jsonify({
                "success": True,
                "message": "Document deleted successfully"
            })
        else:
            return jsonify({
                "success": False,
                "error": "Document reference not found"
            }), 404
            
    except Exception as e:
        logging.error(f"Error deleting task document: {str(e)}")
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500



@app.route('/logout')
def logout():
    session.clear()
    flash('You were successfully logged out')
    return redirect(url_for('login'))

if __name__ == '__main__':
    app.run(debug=True)
